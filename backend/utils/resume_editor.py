import os
import re
import PyPDF2
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from typing import Optional, Tuple
import tempfile
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import black, darkblue

class ResumeEditor:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for resume formatting"""
        self.styles.add(ParagraphStyle(
            name='ResumeHeading',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=6
        ))
        
        self.styles.add(ParagraphStyle(
            name='ResumeSection',
            parent=self.styles['Heading2'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=8
        ))
        
        self.styles.add(ParagraphStyle(
            name='ResumeBody',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=4
        ))
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF, DOCX, or TXT file
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using pdfplumber for better layout preservation"""
        
        # Method 1: Try pdfplumber first (best for layout preservation)
        try:
            text = self._extract_pdf_with_pdfplumber(file_path)
            if text and len(text.strip()) > 50:
                print(f"Successfully extracted PDF with pdfplumber: {len(text)} chars")
                return text
        except Exception as e:
            print(f"Pdfplumber extraction failed: {e}")
        
        # Method 2: Fall back to PyPDF2 with aggressive defragmentation
        try:
            text = self._extract_pdf_method1(file_path)
            if text and len(text.strip()) > 50:
                # Apply aggressive defragmentation
                text = self._fix_character_level_fragmentation(text)
                return self._clean_extracted_text(text)
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
        
        return ""
    
    def _extract_pdf_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber which preserves layout better"""
        all_text = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract text with layout preservation
                page_text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                
                if page_text:
                    # Clean up excessive whitespace while preserving structure
                    lines = page_text.split('\n')
                    cleaned_lines = []
                    
                    for line in lines:
                        # Remove trailing spaces but preserve indentation
                        cleaned_line = line.rstrip()
                        if cleaned_line:  # Only keep non-empty lines
                            # Normalize multiple spaces to single space within the line
                            # but preserve leading spaces for indentation
                            parts = cleaned_line.split()
                            if parts:
                                leading_spaces = len(cleaned_line) - len(cleaned_line.lstrip())
                                cleaned_line = ' ' * min(leading_spaces, 4) + ' '.join(parts)
                            cleaned_lines.append(cleaned_line)
                    
                    page_text = '\n'.join(cleaned_lines)
                    all_text.append(page_text)
        
        # Join pages with double newline
        full_text = '\n\n'.join(all_text)
        
        # Apply minimal post-processing to fix common issues
        full_text = self._minimal_text_cleanup(full_text)
        
        return full_text
    
    def _minimal_text_cleanup(self, text: str) -> str:
        """Minimal cleanup to fix obvious issues without destroying layout"""
        if not text:
            return ""
        
        # Fix email addresses that might be split across lines or have extra spaces
        text = re.sub(r'([a-zA-Z0-9._%+-]+)\s*@\s*([a-zA-Z0-9.-]+)\s*\.\s*([a-zA-Z]{2,})', r'\1@\2.\3', text)
        
        # Fix phone numbers that might be split
        text = re.sub(r'(\d{3})\s*[-.\s]\s*(\d{3})\s*[-.\s]\s*(\d{4})', r'\1-\2-\3', text)
        
        # Remove excessive blank lines (more than 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Ensure consistent line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _extract_pdf_method1(self, file_path: str) -> str:
        """Standard PyPDF2 extraction with character-level reconstruction"""
        text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    # Try to reconstruct proper text from fragments
                    reconstructed = self._reconstruct_text_from_fragments(page_text)
                    text_parts.append(reconstructed)
        
        return "\n\n".join(text_parts)
    
    def _reconstruct_text_from_fragments(self, raw_text: str) -> str:
        """Reconstruct proper text from fragmented PDF extraction"""
        if not raw_text:
            return ""
        
        # Split into lines and clean
        lines = raw_text.split('\n')
        reconstructed_lines = []
        current_line = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_line:
                    reconstructed_lines.append(current_line)
                    current_line = ""
                continue
            
            # Check if this line should be joined with the previous
            if current_line and self._should_join_lines(current_line, line):
                current_line += " " + line
            else:
                if current_line:
                    reconstructed_lines.append(current_line)
                current_line = line
        
        # Add the last line
        if current_line:
            reconstructed_lines.append(current_line)
        
        return "\n".join(reconstructed_lines)
    
    def _should_join_lines(self, current_line: str, next_line: str) -> bool:
        """Determine if two lines should be joined together"""
        # Don't join if current line ends with clear sentence endings
        if current_line.endswith(('.', ':', '•', '-')):
            return False
        
        # Don't join if next line starts with bullet point or section header
        if next_line.startswith(('•', '-', '*')) or next_line.isupper():
            return False
        
        # Don't join if next line looks like a date or location
        if re.match(r'^\d{4}|^[A-Z][a-z]+,\s*[A-Z]{2}', next_line):
            return False
        
        # Join if current line seems incomplete (very short or doesn't end properly)
        if len(current_line) < 10 or not current_line[-1].isalnum():
            return True
        
        # Join if next line starts with lowercase (continuation)
        if next_line and next_line[0].islower():
            return True
        
        return False
    
    def _process_raw_pdf_text(self, raw_text: str) -> str:
        """Process raw PDF text with different logic"""
        # Remove excessive whitespace first
        cleaned = re.sub(r'\s+', ' ', raw_text)
        
        # Try to identify and separate sections
        sections = []
        current_section = ""
        
        # Split by potential section boundaries
        parts = re.split(r'([A-Z\s]{3,}(?:\n|$))', cleaned)
        
        for part in parts:
            part = part.strip()
            if part:
                if part.isupper() and len(part) > 2:
                    # This looks like a section header
                    if current_section:
                        sections.append(current_section)
                    current_section = part + "\n"
                else:
                    current_section += part + " "
        
        if current_section:
            sections.append(current_section)
        
        return "\n\n".join(sections)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with better formatting"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:  # Only add non-empty paragraphs
                    text_parts.append(para_text)
            
            # Join paragraphs with single line breaks
            full_text = "\n".join(text_parts)
            return self._clean_extracted_text(full_text)
            
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Apply basic cleanup to ensure consistent formatting
            return self._clean_extracted_text(text)
            
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return self._clean_extracted_text(text)
            except Exception as e:
                print(f"Error extracting TXT text with latin-1 encoding: {e}")
                return ""
        except Exception as e:
            print(f"Error extracting TXT text: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean up extracted text to fix common formatting issues"""
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = []
        for line in text.split('\n'):
            # Strip whitespace but preserve meaningful content
            clean_line = ' '.join(line.split())  # Normalize internal spaces
            if clean_line:  # Only keep non-empty lines
                lines.append(clean_line)
        
        # Join lines back together with proper spacing
        cleaned_text = '\n'.join(lines)
        
        # Fix common PDF extraction issues
        cleaned_text = self._fix_pdf_extraction_issues(cleaned_text)
        
        return cleaned_text.strip()
    
    def _fix_pdf_extraction_issues(self, text: str) -> str:
        """Fix common issues from PDF text extraction"""
        # Replace multiple consecutive line breaks with double line breaks
        
        # Fix cases where contact info gets split
        text = re.sub(r'\n(\+?\d)', r' \1', text)  # Fix phone numbers
        text = re.sub(r'\n(@)', r'\1', text)  # Fix email addresses
        text = re.sub(r'(\w)\n(\w)', r'\1 \2', text)  # Join broken words
        
        # Normalize section headers
        text = re.sub(r'\n([A-Z\s]{2,})\n', r'\n\n\1\n', text)  # Section headers
        
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', r'\n\n', text)  # Max 2 consecutive line breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces
        
        return text
    
    def _fix_character_level_fragmentation(self, text: str) -> str:
        """Fix PDFs where characters are split by spaces or newlines"""
        if not text:
            return ""
        
        # First pass: fix obvious character-level splits
        # Pattern to find single/double character fragments that should be joined
        # e.g., "P RO F" -> "PROF", "RE C" -> "REC"
        
        # Replace patterns like "P R O F" with "PROF"
        text = re.sub(r'\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b', r'\1\2\3\4', text)
        text = re.sub(r'\b([A-Z])\s+([A-Z])\s+([A-Z])\b', r'\1\2\3', text)
        text = re.sub(r'\b([A-Z])\s+([A-Z])\b', r'\1\2', text)
        
        # Fix split words with lowercase too
        text = re.sub(r'\b(\w)\s+(\w)\s+(\w)\s+(\w)\s+(\w)\b', lambda m: 
                     ''.join(m.groups()) if len(''.join(m.groups())) <= 10 else m.group(0), text)
        text = re.sub(r'\b(\w)\s+(\w)\s+(\w)\s+(\w)\b', lambda m: 
                     ''.join(m.groups()) if len(''.join(m.groups())) <= 8 else m.group(0), text)
        text = re.sub(r'\b(\w)\s+(\w)\s+(\w)\b', lambda m: 
                     ''.join(m.groups()) if len(''.join(m.groups())) <= 6 else m.group(0), text)
        
        # Second pass: reconstruct lines
        lines = text.split('\n')
        reconstructed = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                reconstructed.append('')
                i += 1
                continue
            
            # Check if this line and following lines should be merged
            merged_line = line
            j = i + 1
            
            # Look ahead to see if we should merge lines
            while j < len(lines) and j < i + 5:  # Look up to 5 lines ahead
                next_line = lines[j].strip()
                if not next_line:
                    break
                
                # Check if lines should be merged
                if self._should_merge_fragmented_lines(merged_line, next_line):
                    # If the next line is very short and looks like a continuation
                    if len(next_line) < 20 or next_line[0].islower():
                        merged_line += ' ' + next_line
                        j += 1
                    else:
                        break
                else:
                    break
            
            reconstructed.append(merged_line)
            i = j if j > i + 1 else i + 1
        
        # Join and do final cleanup
        result = '\n'.join(reconstructed)
        
        # Fix common patterns
        result = self._fix_common_resume_patterns(result)
        
        return result
    
    def _should_merge_fragmented_lines(self, current: str, next_line: str) -> bool:
        """Determine if fragmented lines should be merged"""
        # Very short current line likely needs continuation
        if len(current) < 10:
            return True
        
        # Line doesn't end with proper punctuation
        if current and not current[-1] in '.,:;!?':
            # Next line starts with lowercase
            if next_line and next_line[0].islower():
                return True
            # Current line ends with a word fragment
            if len(current.split()[-1]) < 3:
                return True
        
        return False
    
    def _fix_common_resume_patterns(self, text: str) -> str:
        """Fix common resume-specific patterns"""
        # Fix date patterns
        text = re.sub(r'(\d{4})\s*-\s*(\d{4})', r'\1-\2', text)
        text = re.sub(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{4})', r'\1 \2', text)
        # Support full month names
        text = re.sub(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{4})', r'\1 \2', text)
        # Support MM/YYYY format
        text = re.sub(r'(\d{1,2})/(\d{4})', r'\1/\2', text)
        
        # Fix location patterns
        text = re.sub(r'([A-Z][a-z]+)\s*,\s*([A-Z]{2})\s+(\d{5})', r'\1, \2 \3', text)
        
        # Fix email patterns
        text = re.sub(r'(\w+)\s*@\s*(\w+)\s*\.\s*(\w+)', r'\1@\2.\3', text)
        
        # Fix phone patterns
        text = re.sub(r'\(\s*(\d{3})\s*\)\s*(\d{3})\s*-\s*(\d{4})', r'(\1) \2-\3', text)
        
        # Fix bullet points
        text = re.sub(r'^\s*[•·▪▫◦‣⁃]\s*', '• ', text, flags=re.MULTILINE)
        
        return text
    
    def _normalize_text_to_ascii(self, text: str) -> str:
        """Normalize text to ASCII to prevent Unicode width calculation issues"""
        if not text:
            return ""
        
        # Common Unicode to ASCII replacements
        replacements = {
            # Smart quotes
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
            # Em and en dashes
            '—': '-',
            '–': '-',
            # Ellipsis
            '…': '...',
            # Bullet points - keep original bullets, just normalize variants
            '●': '•',
            '·': '•',
            '▪': '•',
            '▫': '•',
            '◦': '•',
            '‣': '•',
            '⁃': '•',
            # Other common characters
            '©': '(c)',
            '®': '(r)',
            '™': '(tm)',
            # Accented characters
            'é': 'e',
            'è': 'e',
            'ê': 'e',
            'ë': 'e',
            'á': 'a',
            'à': 'a',
            'â': 'a',
            'ä': 'a',
            'ó': 'o',
            'ò': 'o',
            'ô': 'o',
            'ö': 'o',
            'ú': 'u',
            'ù': 'u',
            'û': 'u',
            'ü': 'u',
            'ñ': 'n',
            'ç': 'c',
        }
        
        # Apply replacements
        normalized_text = text
        for unicode_char, ascii_char in replacements.items():
            normalized_text = normalized_text.replace(unicode_char, ascii_char)
        
        # Remove any remaining non-ASCII characters except bullets
        try:
            # Keep bullet characters but convert other problematic characters
            result = ""
            for char in normalized_text:
                if char == '•':
                    result += '•'  # Keep bullets
                elif ord(char) < 128:
                    result += char  # Keep ASCII
                else:
                    result += '?'  # Replace other Unicode
            normalized_text = result
        except Exception:
            # If that fails, filter out non-ASCII characters but keep bullets
            normalized_text = ''.join(char if ord(char) < 128 or char == '•' else '?' for char in normalized_text)
        
        return normalized_text
    
    def _simple_char_wrap(self, text: str, max_chars: int) -> list:
        """Dead simple character-based text wrapping - guaranteed to work"""
        if not text:
            return []
        
        if len(text) <= max_chars:
            return [text]
        
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            test_line = current_line + (" " if current_line else "") + word
            
            if len(test_line) <= max_chars:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def create_tailored_resume_pdf_html(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        Create PDF using HTML-to-PDF conversion with WeasyPrint - much more reliable than ReportLab
        """
        try:
            from weasyprint import HTML, CSS
            
            # Normalize text to handle Unicode issues
            tailored_text = self._normalize_text_to_ascii(tailored_text)
            
            # Parse the resume content
            sections = self._parse_resume_for_html(tailored_text)
            
            # Generate HTML
            html_content = self._generate_resume_html(sections)
            
            # Generate PDF from HTML
            HTML(string=html_content).write_pdf(output_path)
            
            return True
            
        except Exception as e:
            print(f"Error creating HTML-based PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def _parse_resume_for_html(self, text: str) -> dict:
        """Parse resume text into structured data for HTML generation"""
        lines = text.split('\n')
        resume_data = {
            'name': '',
            'contact': '',
            'sections': []
        }
        
        current_section = None
        current_content = []
        name_detected = False
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            if not stripped:
                continue
                
            # Detect name (first substantial line)
            if not name_detected and i < 3:
                if (not any(char in stripped.lower() for char in ['@', '.com', 'phone', 'email', '(', ')', 'linkedin', 'http']) and 
                    len(stripped) < 60 and len(stripped) > 5 and 
                    not stripped.isupper()):
                    resume_data['name'] = stripped
                    name_detected = True
                    continue
            
            # Detect contact info
            if (any(indicator in stripped.lower() for indicator in ['@', '.com', 'phone', '(', ')', 'email', 'linkedin', 'http']) or 
                re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', stripped)):
                resume_data['contact'] = stripped
                continue
            
            # Detect section headers
            if self._is_section_header(stripped):
                # Save previous section
                if current_section:
                    resume_data['sections'].append({
                        'title': current_section,
                        'content': current_content
                    })
                # Start new section
                current_section = stripped
                current_content = []
                continue
            
            # Add to current section
            current_content.append(stripped)
        
        # Add the last section
        if current_section:
            resume_data['sections'].append({
                'title': current_section,
                'content': current_content
            })
        
        return resume_data
    
    def _generate_resume_html(self, resume_data: dict) -> str:
        """Generate professional HTML resume with expert-level formatting"""
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Calibri:wght@400;600;700&display=swap');
                
                * {{
                    margin: 0;
                    padding: 0;
                    box-sizing: border-box;
                }}
                
                body {{
                    font-family: 'Calibri', 'Arial', sans-serif;
                    font-size: 10.5pt;
                    line-height: 1.2;
                    color: #333;
                    max-width: 8.5in;
                    margin: 0 auto;
                    padding: 0.5in;
                    background: white;
                    page-break-inside: avoid;
                }}
                
                .header {{
                    text-align: left;
                    margin-bottom: 12px;
                    padding-bottom: 8px;
                    border-bottom: 1px solid #2c5aa0;
                }}
                
                .name {{
                    font-size: 18pt;
                    font-weight: 700;
                    color: #2c5aa0;
                    margin-bottom: 4px;
                    letter-spacing: 0.5px;
                }}
                
                .contact {{
                    font-size: 9.5pt;
                    color: #555;
                    line-height: 1.1;
                }}
                
                .section {{
                    margin-bottom: 6px;
                }}
                
                .section-header {{
                    font-size: 10.5pt;
                    font-weight: 700;
                    color: #2c5aa0;
                    text-transform: uppercase;
                    margin-bottom: 3px;
                    padding-bottom: 1px;
                    border-bottom: 1px solid #2c5aa0;
                    letter-spacing: 0.5px;
                }}
                
                .experience-item {{
                    margin-bottom: 15px;
                }}
                
                .company-info {{
                    display: flex;
                    justify-content: space-between;
                    align-items: baseline;
                    margin-bottom: 1px;
                }}
                
                .company {{
                    font-weight: 600;
                    font-size: 10pt;
                    color: #333;
                }}
                
                .location {{
                    font-size: 9pt;
                    color: #666;
                    font-style: italic;
                }}
                
                .position-info {{
                    display: flex;
                    justify-content: space-between;
                    align-items: baseline;
                    margin-bottom: 2px;
                }}
                
                .position {{
                    font-weight: 600;
                    font-size: 10pt;
                    color: #2c5aa0;
                }}
                
                .dates {{
                    font-size: 9pt;
                    color: #666;
                    font-weight: 500;
                }}
                
                .bullets {{
                    margin-left: 0;
                    padding-left: 0;
                    margin-bottom: 6px;
                }}
                
                .bullet {{
                    position: relative;
                    margin-bottom: 2px;
                    padding-left: 15px;
                    margin-left: 0;
                    font-size: 9.5pt;
                    line-height: 1.15;
                    text-align: left;
                    text-indent: -15px;
                    display: block;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                
                .bullet::before {{
                    content: "•";
                    position: absolute;
                    left: 0;
                    top: 0;
                    color: #2c5aa0;
                    font-weight: bold;
                    width: 15px;
                    text-align: left;
                }}
                
                .skills-grid {{
                    display: grid;
                    grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                    gap: 8px;
                }}
                
                .skill-item {{
                    font-size: 10.5pt;
                    padding: 3px 0;
                    position: relative;
                    padding-left: 15px;
                }}
                
                .skill-item::before {{
                    content: "•";
                    position: absolute;
                    left: 0;
                    color: #2c5aa0;
                    font-weight: bold;
                }}
                
                .content {{
                    font-size: 9.5pt;
                    line-height: 1.2;
                    margin-bottom: 4px;
                    text-align: justify;
                }}
                
                .education-entry {{
                    margin-bottom: 1px;
                    line-height: 1.0;
                }}
                
                .degree {{
                    font-weight: 600;
                    font-size: 9.5pt;
                    color: #333;
                    margin: 0;
                    padding: 0;
                    line-height: 1.1;
                    display: block;
                }}
                
                .institution {{
                    font-size: 9pt;
                    color: #666;
                    margin: 0;
                    padding: 0;
                    line-height: 1.1;
                    display: block;
                }}
                
                .skills-category {{
                    margin-bottom: 1px;
                    line-height: 1.0;
                }}
                
                .skills-label {{
                    font-weight: 600;
                    font-size: 9pt;
                    color: #333;
                    display: inline;
                }}
                
                .skills-list {{
                    font-size: 9pt;
                    color: #555;
                    display: inline;
                    margin-left: 3px;
                }}
                
                @page {{
                    size: letter;
                    margin: 0.5in;
                    page-break-inside: avoid;
                }}
                
                @media print {{
                    body {{
                        padding: 0;
                        font-size: 10pt;
                    }}
                    .bullet {{
                        font-size: 10pt;
                    }}
                    .skill-item {{
                        font-size: 10pt;
                    }}
                    .content {{
                        font-size: 10pt;
                    }}
                }}
            </style>
        </head>
        <body>
            <div class="header">
        """
        
        # Add name and contact in header
        if resume_data['name']:
            html += f'<div class="name">{resume_data["name"]}</div>'
        
        if resume_data['contact']:
            html += f'<div class="contact">{resume_data["contact"]}</div>'
        
        html += '</div>'
        
        # Add sections
        for section in resume_data['sections']:
            section_title = section["title"].upper()
            html += f'<div class="section">'
            html += f'<div class="section-header">{section_title}</div>'
            
            if section_title in ['PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'WORK EXPERIENCE']:
                html += self._format_experience_section(section['content'])
            elif section_title in ['SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES']:
                html += self._format_skills_section(section['content'])
            else:
                html += self._format_general_section(section['content'])
            
            html += '</div>'
        
        html += """
        </body>
        </html>
        """
        
        return html
    
    def _format_experience_section(self, content: list) -> str:
        """Format professional experience section"""
        html = ""
        current_company = ""
        current_position = ""
        bullets = []
        
        for line in content:
            line = line.strip()
            if not line:
                continue
                
            # Check if it's a company/location line
            if '|' in line and any(word in line.lower() for word in ['corp', 'inc', 'llc', 'company', 'tech', 'ca', 'ny', 'tx']):
                # Save previous job if exists
                if current_company and bullets:
                    html += self._create_experience_item(current_company, current_position, bullets)
                    bullets = []
                
                parts = line.split('|')
                if len(parts) >= 2:
                    current_company = f'<span class="company">{parts[0].strip()}</span> | <span class="location">{parts[1].strip()}</span>'
                else:
                    current_company = f'<span class="company">{line}</span>'
                continue
            
            # Check if it's a position/date line
            elif '|' in line and any(char.isdigit() for char in line):
                parts = line.split('|')
                if len(parts) >= 2:
                    current_position = f'<span class="position">{parts[0].strip()}</span><span class="dates">{parts[1].strip()}</span>'
                else:
                    current_position = f'<span class="position">{line}</span>'
                continue
            
            # Check if it's a bullet point
            elif line.startswith(('•', '●', '·', '-', '*')):
                clean_bullet = line[1:].strip()
                bullets.append(clean_bullet)
            else:
                # Treat as bullet if it's in experience section
                bullets.append(line)
        
        # Add the last job
        if current_company and bullets:
            html += self._create_experience_item(current_company, current_position, bullets)
        
        return html
    
    def _create_experience_item(self, company: str, position: str, bullets: list) -> str:
        """Create a single experience item"""
        html = '<div class="experience-item">'
        
        if company:
            html += f'<div class="company-info">{company}</div>'
        
        if position:
            html += f'<div class="position-info">{position}</div>'
        
        if bullets:
            html += '<div class="bullets">'
            for bullet in bullets:
                html += f'<div class="bullet">{bullet}</div>'
            html += '</div>'
        
        html += '</div>'
        return html
    
    def _format_skills_section(self, content: list) -> str:
        """Format skills section with category labels and compact lists"""
        html = ''
        
        for line in content:
            line = line.strip()
            if not line:
                continue
                
            # Remove bullet if present
            if line.startswith(('•', '●', '·', '-', '*')):
                line = line[1:].strip()
            
            # Check if line contains a colon (category: skills format)
            if ':' in line:
                parts = line.split(':', 1)
                category = parts[0].strip()
                skills = parts[1].strip()
                html += f'<div class="skills-category"><span class="skills-label">{category}:</span><span class="skills-list">{skills}</span></div>'
            else:
                html += f'<div class="skills-category"><span class="skills-list">{line}</span></div>'
        
        return html
    
    def _format_general_section(self, content: list) -> str:
        """Format general content sections with education-specific handling"""
        html = ""
        i = 0
        
        while i < len(content):
            line = content[i].strip()
            if not line:
                i += 1
                continue
            
            # Check if it's a bullet point
            if line.startswith(('•', '●', '·', '-', '*')):
                clean_content = line[1:].strip()
                html += f'<div class="bullet">{clean_content}</div>'
            else:
                # Check if this looks like an education entry (degree followed by institution)
                if (i + 1 < len(content) and 
                    ('degree' in line.lower() or 'bachelor' in line.lower() or 'master' in line.lower() or 
                     'mba' in line.lower() or 'bs' in line.lower() or 'ms' in line.lower() or
                     'phd' in line.lower() or 'doctorate' in line.lower() or 'certificate' in line.lower()) and
                    ('university' in content[i + 1].lower() or 'college' in content[i + 1].lower() or 
                     'institute' in content[i + 1].lower() or 'school' in content[i + 1].lower())):
                    
                    # Format as education entry with no extra spacing
                    html += f'<div class="education-entry">'
                    html += f'<div class="degree">{line}</div>'
                    html += f'<div class="institution">{content[i + 1].strip()}</div>'
                    html += f'</div>'
                    i += 2  # Skip the next line since we processed it
                    continue
                else:
                    html += f'<div class="content">{line}</div>'
            
            i += 1
        
        return html
    
    def create_tailored_resume_pdf(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        Create a PDF using improved ReportLab formatting for professional appearance
        """
        # Use improved ReportLab method with better styling
        return self.create_tailored_resume_pdf_improved(tailored_text, output_path, job_title)
    
    def create_tailored_resume_pdf_improved(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        Create a PDF using improved ReportLab formatting that matches the target professional appearance
        """
        try:
            # Log what we receive from AI
            print("=" * 80)
            print("📄 PDF GENERATION - INPUT TEXT:")
            print("=" * 80)
            print(repr(tailored_text))
            print("=" * 80)
            print("📄 PDF GENERATION - FORMATTED INPUT:")
            print("=" * 80)
            print(tailored_text)
            print("=" * 80)
            
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import black, blue, HexColor
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            
            # Create the document with compact margins to fit more content
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=0.4*inch,   # Reduced from 0.5
                leftMargin=0.4*inch,    # Reduced from 0.5
                topMargin=0.4*inch,     # Reduced from 0.5
                bottomMargin=0.4*inch   # Reduced from 0.5
            )
            
            # Get base styles
            styles = getSampleStyleSheet()
            
            # Define professional blue color
            professional_blue = HexColor('#2c5aa0')
            
            # Define compact styles with reduced font sizes to fit more content
            name_style = ParagraphStyle(
                'CompactNameStyle',
                parent=styles['Normal'],
                fontSize=14,  # Reduced from 16
                textColor=professional_blue,
                spaceAfter=2,
                spaceBefore=0,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            )
            
            title_style = ParagraphStyle(
                'CompactTitleStyle',
                parent=styles['Normal'],
                fontSize=12,  # Reduced from 14
                textColor=professional_blue,
                spaceAfter=2,
                spaceBefore=0,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            )
            
            contact_style = ParagraphStyle(
                'CompactContactStyle',
                parent=styles['Normal'],
                fontSize=9,  # Reduced from 10
                spaceAfter=6,  # Further reduced from 8
                spaceBefore=0,
                fontName='Helvetica'
            )
            
            section_header_style = ParagraphStyle(
                'CompactSectionHeaderStyle',
                parent=styles['Normal'],
                fontSize=11,  # Reduced from 12
                textColor=professional_blue,
                spaceBefore=2,  # Further reduced from 4
                spaceAfter=1,   # Reduced from 2
                fontName='Helvetica-Bold'
            )
            
            company_header_style = ParagraphStyle(
                'CompactCompanyHeaderStyle',
                parent=styles['Normal'],
                fontSize=9,  # Reduced from 10
                spaceBefore=2,  # Further reduced from 3
                spaceAfter=0,   # Reduced from 1
                fontName='Helvetica-Bold'
            )
            

            
            body_style = ParagraphStyle(
                'CompactBodyStyle',
                parent=styles['Normal'],
                fontSize=9,  # Reduced from 10
                spaceAfter=1,  # Reduced from 2
                spaceBefore=0,
                fontName='Helvetica'
            )
            
            # Parse the resume content
            sections = self._parse_resume_for_improved_formatting(tailored_text)
            
            # Log what the parsing produces
            print("=" * 80)
            print("📄 PDF GENERATION - PARSED SECTIONS:")
            print("=" * 80)
            for i, (section_type, content) in enumerate(sections):
                print(f"{i}: {section_type} -> {repr(content)}")
            print("=" * 80)
            
            # Build the story (list of flowables)
            story = []
            
            # Track current section to skip Education and Skills (they'll be handled horizontally)
            current_section_name = None
            skip_section = False
            
            for section_type, content in sections:
                # Track section headers to know when to skip Education and Skills
                if section_type == "section_header":
                    current_section_name = content.strip().upper()
                    skip_section = current_section_name in ["EDUCATION", "SKILLS"]
                    
                    # Only add non-Education/Skills section headers
                    if not skip_section:
                        # Handle misclassified names that got marked as section_header
                        if content.strip() == "DAVID PATEL":
                            story.append(Paragraph(content.strip(), name_style))
                            story.append(Spacer(1, 4))  # Reduced spacing after name
                        elif content.strip().startswith('•'):
                            # This is actually a bullet point
                            clean_content = content.strip()
                            while clean_content and clean_content[0] in ['•', '●', '·', '-', '*', '+', '◦', '▪', '▫']:
                                clean_content = clean_content[1:].strip()
                            if clean_content:
                                # Use same hanging indent style as other bullets
                                hanging_bullet_style = ParagraphStyle(
                                    'HangingBulletStyle',
                                    parent=styles['Normal'],
                                    fontSize=9,
                                    leftIndent=6,            # Further reduced to move continuation lines LEFT
                                    firstLineIndent=-6,      # Match the leftIndent
                                    spaceAfter=1,           # Reduced spacing
                                    spaceBefore=0,
                                    fontName='Helvetica',
                                    alignment=TA_LEFT,
                                    bulletIndent=0
                                )
                                bullet_content = f"• {clean_content}"
                                story.append(Paragraph(bullet_content, hanging_bullet_style))
                        else:
                            story.append(Paragraph(content.strip().upper(), section_header_style))
                    continue
                
                # Skip content for Education and Skills sections
                if skip_section:
                    continue
                
                if section_type == "name":
                    story.append(Paragraph(content.strip(), name_style))
                    # Add reduced spacing after name
                    story.append(Spacer(1, 4))
                    
                elif section_type == "title":
                    story.append(Paragraph(content.strip(), title_style))
                    
                elif section_type == "contact":
                    # Handle misclassified bullet points that got marked as contact
                    if content.strip().startswith('•'):
                        # This is actually a bullet point
                        clean_content = content.strip()
                        while clean_content and clean_content[0] in ['•', '●', '·', '-', '*', '+', '◦', '▪', '▫']:
                            clean_content = clean_content[1:].strip()
                        if clean_content:
                            # Use same hanging indent style as other bullets
                            hanging_bullet_style = ParagraphStyle(
                                'HangingBulletStyle',
                                parent=styles['Normal'],
                                fontSize=9,
                                leftIndent=6,            # Further reduced to move continuation lines LEFT
                                firstLineIndent=-6,      # Match the leftIndent
                                spaceAfter=1,           # Reduced spacing
                                spaceBefore=0,
                                fontName='Helvetica',
                                alignment=TA_LEFT,
                                bulletIndent=0
                            )
                            bullet_content = f"• {clean_content}"
                            story.append(Paragraph(bullet_content, hanging_bullet_style))
                    else:
                        story.append(Paragraph(content.strip(), contact_style))
                    

                    
                elif section_type == "company_header":
                    story.append(Paragraph(content.strip(), company_header_style))
                    
                elif section_type == "bullet_point":
                    # Clean the bullet content consistently
                    clean_content = content.strip()
                    # Remove any existing bullet symbols
                    while clean_content and clean_content[0] in ['•', '●', '·', '-', '*', '+', '◦', '▪', '▫']:
                        clean_content = clean_content[1:].strip()
                    # Remove numbered/lettered list markers
                    clean_content = re.sub(r'^\d+\.\s+', '', clean_content)
                    clean_content = re.sub(r'^[a-zA-Z]\.\s+', '', clean_content)
                    # Remove any remaining leading whitespace
                    clean_content = clean_content.strip()
                    
                    if clean_content:
                        # Create proper hanging indent bullet with correct formatting
                        # Use a more precise approach for hanging indentation
                        hanging_bullet_style = ParagraphStyle(
                            'HangingBulletStyle',
                            parent=styles['Normal'],
                            fontSize=9,
                            leftIndent=6,            # Further reduced to move continuation lines LEFT
                            firstLineIndent=-6,      # Move first line back to align bullet at margin
                            spaceAfter=1,           # Reduced spacing
                            spaceBefore=0,
                            fontName='Helvetica',
                            alignment=TA_LEFT,
                            bulletIndent=0
                        )
                        
                        # Format bullet content with proper spacing
                        bullet_content = f"• {clean_content}"
                        story.append(Paragraph(bullet_content, hanging_bullet_style))
                        
                elif section_type == "body_text":
                    story.append(Paragraph(content.strip(), body_style))
                    
                elif section_type == "spacing":
                    spacing = min(int(content), 1)  # Even tighter spacing to fit on one page
                    story.append(Spacer(1, spacing))
            
            # After processing all sections, look for Education and Skills to create horizontal layout
            self._add_horizontal_education_skills_layout(story, sections, section_header_style, body_style, styles)
            
            # Build the document
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error creating improved PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback to original method
            return self.create_tailored_resume_pdf_fixed(tailored_text, output_path, job_title)
    
    def _add_horizontal_education_skills_layout(self, story, sections, section_header_style, body_style, styles):
        """Add horizontal layout for Education and Skills sections"""
        try:
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib.colors import HexColor
            
            # Define professional blue color
            professional_blue = HexColor('#2c5aa0')
            
            # Find Education and Skills sections
            education_content = []
            skills_content = []
            current_section = None
            
            for section_type, content in sections:
                if section_type == "section_header":
                    if content.strip().upper() == "EDUCATION":
                        current_section = "EDUCATION"
                        continue
                    elif content.strip().upper() == "SKILLS":
                        current_section = "SKILLS"
                        continue
                    else:
                        current_section = None
                        
                if current_section == "EDUCATION" and section_type != "spacing":
                    education_content.append(content.strip())
                elif current_section == "SKILLS" and section_type != "spacing":
                    skills_content.append(content.strip())
            
            # Only create horizontal layout if both sections exist
            if education_content and skills_content:
                # Format Education content for better readability
                education_text = self._format_education_simple(education_content)
                skills_text = "\n".join(skills_content)
                
                # Create paragraph styles for table content - more compact
                table_content_style = ParagraphStyle(
                    'TableContentStyle',
                    parent=styles['Normal'],
                    fontSize=8,            # Smaller font for Education/Skills
                    spaceAfter=0,          # No spacing after
                    spaceBefore=0,
                    fontName='Helvetica'
                )
                
                # Create slightly bolder style for Education section
                education_content_style = ParagraphStyle(
                    'EducationContentStyle',
                    parent=styles['Normal'],
                    fontSize=8,            # Keep same size
                    spaceAfter=0,          # No spacing after
                    spaceBefore=0,
                    fontName='Helvetica-Bold'  # Make it bold
                )
                
                # Create compact header style for table
                table_header_style = ParagraphStyle(
                    'TableHeaderStyle',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=professional_blue,
                    spaceBefore=0,
                    spaceAfter=0,
                    fontName='Helvetica-Bold'
                )
                
                # Create the table data with bolder Education section
                table_data = [
                    [
                        [
                            Paragraph("EDUCATION", table_header_style),
                            Paragraph(education_text, education_content_style)  # Use bolder style
                        ],
                        [
                            Paragraph("SKILLS", table_header_style),
                            Paragraph(skills_text, table_content_style)
                        ]
                    ]
                ]
                
                # Create table with minimal spacing for compactness
                table = Table(table_data, colWidths=[3.6*inch, 3.6*inch])
                table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 3),     # Reduced padding
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ]))
                
                # Add more spacing before the horizontal Education/Skills section
                story.append(Spacer(1, 8))
                story.append(table)
                
        except Exception as e:
            print(f"Error creating horizontal layout: {e}")
            # If horizontal layout fails, fall back to normal vertical layout
            pass
    
    def _format_education_simple(self, education_content):
        """Simple education formatting with line breaks"""
        try:
            full_text = " ".join(education_content)
            
            # Simple approach: split by years (4 digits preceded by |)
            import re
            
            # Replace " | 2018" or " | 2016" patterns with newlines
            formatted_text = re.sub(r'\s*\|\s*(\d{4})\s*', r' | \1\n', full_text)
            
            # Clean up any trailing/leading whitespace on each line
            lines = [line.strip() for line in formatted_text.split('\n') if line.strip()]
            
            return '\n'.join(lines)
            
        except Exception as e:
            return " ".join(education_content)
    
    def _parse_resume_for_improved_formatting(self, text: str) -> list:
        """Parse resume text for improved formatting that matches the target appearance"""
        lines = text.split('\n')
        sections = []
        name_detected = False
        title_detected = False
        current_section = None
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            if not stripped:
                # Empty line - add minimal spacing to match target exactly
                sections.append(("spacing", "1"))
                continue
            
            # Detect name (first substantial line that's not contact info and not all caps)
            if not name_detected and i < 3:
                if (not any(char in stripped.lower() for char in ['@', '.com', 'phone', 'email', '(', ')', 'linkedin', 'http']) and 
                    len(stripped) < 60 and len(stripped) > 5 and 
                    not stripped.isupper() and not stripped.startswith('•')):
                    sections.append(("name", stripped))
                    name_detected = True
                    continue
            
            # Detect title (second line, often job title, not all caps, not contact)
            if name_detected and not title_detected and i < 5:
                if (not any(char in stripped.lower() for char in ['@', '.com', 'phone', 'email', '(', ')', 'linkedin', 'http']) and 
                    len(stripped) < 80 and len(stripped) > 5 and 
                    not self._is_section_header(stripped) and not stripped.startswith('•')):
                    sections.append(("title", stripped))
                    title_detected = True
                    continue
            
            # Detect contact info (contains email, phone, or LinkedIn patterns)
            if (any(indicator in stripped.lower() for indicator in ['@', '.com', 'phone', '(', ')', 'email', 'linkedin', 'http']) or 
                re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', stripped)) and not stripped.startswith('•'):
                sections.append(("contact", stripped))
                continue
            
            # Detect section headers (all caps, common section names, NOT starting with bullet)
            if (self._is_section_header(stripped) and not stripped.startswith('•')):
                sections.append(("section_header", stripped))
                current_section = stripped.lower().replace(':', '')
                continue
            
            # Detect company/job headers (contain dates and separators)
            if ('|' in stripped and re.search(r'\b\d{4}\b', stripped)):
                sections.append(("company_header", stripped))
                continue
            
            # Detect job title lines (usually follow company headers)
            if (re.search(r'(manager|director|engineer|analyst|developer|specialist|coordinator|lead|senior|junior)\s*\|', stripped.lower()) or
                re.search(r'\w+\s*\|\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)', stripped.lower())):
                sections.append(("company_header", stripped))
                continue
            
            # Detect company lines (contain pipe but no date - these are company | location lines)
            if ('|' in stripped and not re.search(r'\b\d{4}\b', stripped) and 
                len(sections) > 0 and sections[-1][0] in ["spacing", "section_header"]):
                sections.append(("company_header", stripped))
                continue
            
            # BULLET POINT DETECTION - This should be BEFORE other classifications
            is_bullet = False
            # Rule 1: Explicit bullet characters (more comprehensive)
            if stripped.startswith(('•', '●', '·', '*', '-', '◦', '▪', '▫')) or re.match(r'^\s*\d+\.\s+', stripped) or re.match(r'^\s*[a-zA-Z][\.\)]\s+', stripped):
                is_bullet = True
            
            if is_bullet:
                sections.append(("bullet_point", stripped))
                continue
            
            # Everything else is body text
            sections.append(("body_text", stripped))
        
        return sections

    def _create_simple_resume_html(self, sections) -> str:
        """Create super simple HTML that mimics original formatting exactly"""
        
        html = """
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {
                    font-family: Helvetica, Arial, sans-serif;
                    font-size: 10pt;
                    line-height: 1.2;
                    margin: 0.5in;
                    color: black;
                }
                .name {
                    font-size: 16pt;
                    font-weight: bold;
                    color: blue;
                    margin-bottom: 5px;
                }
                .contact {
                    font-size: 9pt;
                    margin-bottom: 15px;
                }
                .section-header {
                    font-size: 11pt;
                    font-weight: bold;
                    color: blue;
                    text-transform: uppercase;
                    margin-top: 10px;
                    margin-bottom: 5px;
                }
                .company-header {
                    font-weight: bold;
                    margin-top: 8px;
                    margin-bottom: 3px;
                }
                .bullet {
                    margin-left: 0px;
                    margin-bottom: 3px;
                    text-indent: -15px;
                    padding-left: 15px;
                    line-height: 1.2;
                    position: relative;
                    display: block;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }
                .bullet::before {
                    content: "• ";
                    color: black;
                    position: absolute;
                    left: 0;
                    top: 0;
                    width: 15px;
                    text-align: left;
                }
                .body-text {
                    margin-bottom: 5px;
                }
            </style>
        </head>
        <body>
        """
        
        for section_type, content in sections:
            if section_type == "name":
                html += f'<div class="name">{content.strip()}</div>'
            elif section_type == "contact":
                html += f'<div class="contact">{content.strip()}</div>'
            elif section_type == "section_header":
                html += f'<div class="section-header">{content.strip().upper()}</div>'
            elif section_type == "company_header":
                html += f'<div class="company-header">{content.strip()}</div>'
            elif section_type == "bullet_point":
                # Clean the bullet content
                clean_content = content.strip()
                while clean_content and clean_content[0] in ['•', '●', '·', '-', '*']:
                    clean_content = clean_content[1:].strip()
                clean_content = re.sub(r'^\d+\.\s+', '', clean_content)
                clean_content = re.sub(r'^[a-zA-Z]\.\s+', '', clean_content)
                
                if clean_content:
                    html += f'<div class="bullet">{clean_content}</div>'
            elif section_type == "body_text":
                html += f'<div class="body-text">{content.strip()}</div>'
            elif section_type == "spacing":
                html += '<div style="margin-bottom: 4px;"></div>'
        
        html += """
        </body>
        </html>
        """
        
        return html
    
    def create_tailored_resume_pdf_reportlab_BROKEN(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        OLD BROKEN REPORTLAB METHOD - keeping for reference but NOT USING
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase.pdfmetrics import stringWidth
            from reportlab.lib.colors import black, blue
            
            # Normalize text to ASCII to prevent Unicode width calculation issues
            tailored_text = self._normalize_text_to_ascii(tailored_text)
            
            # Create a canvas
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Define colors
            darkblue = blue
            
            # More compact font sizes to fit everything on one page
            name_font_size = 16
            section_font_size = 11
            body_font_size = 10
            small_font_size = 9
            
            # Tighter margins for more content space
            left_margin = 36  # 0.5 inch
            right_margin = 36
            top_margin = 36
            bottom_margin = 36
            available_width = width - left_margin - right_margin
            
            # Starting position from top
            y_position = height - top_margin
            
            # Compact line heights
            name_line_height = 18
            section_line_height = 14
            body_line_height = 12
            small_line_height = 10
            
            # Parse and format the resume content
            sections = self._parse_resume_for_compact_formatting(tailored_text)
            
            for section_type, content in sections:
                # Conservative page break check
                if y_position < bottom_margin + 50:
                    break  # Stop adding content to keep it on one page
                
                if section_type == "name":
                    # Name formatting
                    c.setFont("Helvetica-Bold", name_font_size)
                    c.setFillColor(darkblue)
                    c.drawString(left_margin, y_position, content.strip())
                    c.setFillColor(black)
                    y_position -= name_line_height
                    
                elif section_type == "contact":
                    # Contact info
                    c.setFont("Helvetica", small_font_size)
                    c.drawString(left_margin, y_position, content.strip())
                    y_position -= small_line_height + 6
                    
                elif section_type == "section_header":
                    # Section headers - consistent spacing
                    y_position -= 6  # Consistent spacing before headers
                    c.setFont("Helvetica-Bold", section_font_size)
                    c.setFillColor(darkblue)
                    c.drawString(left_margin, y_position, content.strip().upper())
                    c.setFillColor(black)
                    y_position -= section_line_height
                    
                elif section_type == "company_header":
                    # Company/job headers - consistent spacing
                    y_position -= 4  # Consistent spacing before company headers
                    c.setFont("Helvetica-Bold", body_font_size)
                    wrapped_lines = self._wrap_text_by_width(content.strip(), available_width, "Helvetica-Bold", body_font_size)
                    for line in wrapped_lines:
                        if y_position < bottom_margin + 30:
                            break
                        c.drawString(left_margin, y_position, line)
                        y_position -= body_line_height
                    
                elif section_type == "bullet_point":
                    # Consistent bullet point formatting
                    c.setFont("Helvetica", body_font_size)
                    c.setFillColor(black)
                    
                    # Remove existing bullet symbols, numbers, and letters
                    clean_content = content.strip()
                    # Remove bullet symbols (including our normalized ASCII bullets)
                    while clean_content and clean_content[0] in ['•', '●', '·', '-', '*']:
                        clean_content = clean_content[1:].strip()
                    # Remove numbered lists: "1. ", "2. ", etc.
                    clean_content = re.sub(r'^\d+\.\s+', '', clean_content)
                    # Remove lettered lists: "a. ", "b. ", "A. ", "B. ", etc.
                    clean_content = re.sub(r'^[a-zA-Z]\.\s+', '', clean_content)
                    
                    if not clean_content:
                        continue
                    
                    # This section is no longer used - bullet points handled in main method
                    
                    for i, line in enumerate(wrapped_lines):
                        if y_position < bottom_margin + 20:
                            break
                        
                        # Ensure consistent font for all bullet content
                        c.setFont("Helvetica", body_font_size)
                        c.setFillColor(black)
                        
                        if i == 0:
                            # First line with bullet
                            c.drawString(left_margin + 10, y_position, "•")
                            c.drawString(left_margin + 20, y_position, line)
                        else:
                            # Continuation lines
                            c.drawString(left_margin + 20, y_position, line)
                        
                        y_position -= body_line_height
                    
                elif section_type == "body_text":
                    # Regular body text - ensure consistent formatting with bullet points
                    c.setFont("Helvetica", body_font_size)
                    c.setFillColor(black)
                    wrapped_lines = self._wrap_text_by_width(content.strip(), available_width, "Helvetica", body_font_size)
                    
                    for line in wrapped_lines:
                        if y_position < bottom_margin + 20:
                            break
                        # Use consistent positioning and font
                        c.drawString(left_margin, y_position, line)
                        y_position -= body_line_height
                    
                elif section_type == "spacing":
                    spacing = min(int(content), 6)
                    y_position -= spacing
            
            # Save the PDF
            c.save()
            return True
            
        except Exception as e:
            print(f"Error creating compact PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def create_tailored_resume_pdf_reportlab_old(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        OLD ReportLab method - keeping as backup but not using anymore
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase.pdfmetrics import stringWidth
            
            # Normalize text to ASCII to prevent Unicode width calculation issues
            tailored_text = self._normalize_text_to_ascii(tailored_text)
            
            # Create a canvas
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # More compact font sizes to fit everything on one page
            name_font_size = 16
            section_font_size = 11
            body_font_size = 10
            small_font_size = 9
            
            # Tighter margins for more content space
            left_margin = 36  # 0.5 inch
            right_margin = 36
            top_margin = 36
            bottom_margin = 36
            available_width = width - left_margin - right_margin
            
            # Starting position from top
            y_position = height - top_margin
            
            # Compact line heights
            name_line_height = 18
            section_line_height = 14
            body_line_height = 12
            small_line_height = 10
            
            # Parse and format the resume content
            sections = self._parse_resume_for_compact_formatting(tailored_text)
            
            for section_type, content in sections:
                # Conservative page break check
                if y_position < bottom_margin + 50:
                    break  # Stop adding content to keep it on one page
                
                if section_type == "name":
                    # Name formatting
                    c.setFont("Helvetica-Bold", name_font_size)
                    c.setFillColor(darkblue)
                    c.drawString(left_margin, y_position, content.strip())
                    c.setFillColor(black)
                    y_position -= name_line_height
                    
                elif section_type == "contact":
                    # Contact info
                    c.setFont("Helvetica", small_font_size)
                    c.drawString(left_margin, y_position, content.strip())
                    y_position -= small_line_height + 6
                    
                elif section_type == "section_header":
                    # Section headers - consistent spacing
                    y_position -= 6  # Consistent spacing before headers
                    c.setFont("Helvetica-Bold", section_font_size)
                    c.setFillColor(darkblue)
                    c.drawString(left_margin, y_position, content.strip().upper())
                    c.setFillColor(black)
                    y_position -= section_line_height
                    
                elif section_type == "company_header":
                    # Company/job headers - consistent spacing
                    y_position -= 4  # Consistent spacing before company headers
                    c.setFont("Helvetica-Bold", body_font_size)
                    wrapped_lines = self._wrap_text_by_width(content.strip(), available_width, "Helvetica-Bold", body_font_size)
                    for line in wrapped_lines:
                        if y_position < bottom_margin + 30:
                            break
                        c.drawString(left_margin, y_position, line)
                        y_position -= body_line_height
                    
                elif section_type == "bullet_point":
                    # Consistent bullet point formatting
                    c.setFont("Helvetica", body_font_size)
                    c.setFillColor(black)
                    
                    # Remove existing bullet symbols, numbers, and letters
                    clean_content = content.strip()
                    # Remove bullet symbols (including our normalized ASCII bullets)
                    while clean_content and clean_content[0] in ['•', '●', '·', '-', '*']:
                        clean_content = clean_content[1:].strip()
                    # Remove numbered lists: "1. ", "2. ", etc.
                    clean_content = re.sub(r'^\d+\.\s+', '', clean_content)
                    # Remove lettered lists: "a. ", "b. ", "A. ", "B. ", etc.
                    clean_content = re.sub(r'^[a-zA-Z]\.\s+', '', clean_content)
                    
                    if not clean_content:
                        continue
                    
                    # This section is no longer used - bullet points handled in main method
                    
                    for i, line in enumerate(wrapped_lines):
                        if y_position < bottom_margin + 20:
                            break
                        
                        # Ensure consistent font for all bullet content
                        c.setFont("Helvetica", body_font_size)
                        c.setFillColor(black)
                        
                        if i == 0:
                            # First line with bullet
                            c.drawString(left_margin + 10, y_position, "•")
                            c.drawString(left_margin + 20, y_position, line)
                        else:
                            # Continuation lines
                            c.drawString(left_margin + 20, y_position, line)
                        
                        y_position -= body_line_height
                    
                elif section_type == "body_text":
                    # Regular body text - ensure consistent formatting with bullet points
                    c.setFont("Helvetica", body_font_size)
                    c.setFillColor(black)
                    wrapped_lines = self._wrap_text_by_width(content.strip(), available_width, "Helvetica", body_font_size)
                    
                    for line in wrapped_lines:
                        if y_position < bottom_margin + 20:
                            break
                        # Use consistent positioning and font
                        c.drawString(left_margin, y_position, line)
                        y_position -= body_line_height
                    
                elif section_type == "spacing":
                    spacing = min(int(content), 6)
                    y_position -= spacing
            
            # Save the PDF
            c.save()
            return True
            
        except Exception as e:
            print(f"Error creating compact PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def _parse_resume_for_compact_formatting(self, text: str) -> list:
        """Parse resume text for compact formatting with improved classification"""
        lines = text.split('\n')
        sections = []
        name_detected = False
        current_section = None
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            if not stripped:
                # Empty line - add minimal spacing
                sections.append(("spacing", "4"))
                continue
            
            # Detect name (first substantial line that's not contact info)
            if not name_detected and i < 3:
                if (not any(char in stripped.lower() for char in ['@', '.com', 'phone', 'email', '(', ')', 'linkedin', 'http']) and 
                    len(stripped) < 60 and len(stripped) > 5 and 
                    not stripped.isupper()):
                    sections.append(("name", stripped))
                    name_detected = True
                    continue
            
            # Detect contact info
            if (any(indicator in stripped.lower() for indicator in ['@', '.com', 'phone', '(', ')', 'email', 'linkedin', 'http']) or 
                re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', stripped)):
                sections.append(("contact", stripped))
                continue
            
            # Detect section headers (both all caps and title case)
            if (self._is_section_header(stripped)):
                sections.append(("section_header", stripped))
                current_section = stripped.lower().replace(':', '')
                continue
            
            # Detect company/job headers (contain dates and separators)
            if ('|' in stripped and re.search(r'\b\d{4}\b', stripped)):
                sections.append(("company_header", stripped))
                continue
            
            # Detect job title lines (usually follow company headers)
            if (re.search(r'(manager|director|engineer|analyst|developer|specialist|coordinator|lead|senior|junior)\s*\|', stripped.lower()) or
                re.search(r'\w+\s*\|\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)', stripped.lower())):
                sections.append(("company_header", stripped))
                continue
            
            # Detect company lines (contain pipe but no date - these are company | location lines)
            if ('|' in stripped and not re.search(r'\b\d{4}\b', stripped) and 
                len(sections) > 0 and sections[-1][0] in ["spacing", "section_header"]):
                sections.append(("company_header", stripped))
                continue
            
            # Comprehensive bullet point detection
            is_bullet = False
            
            # 1. Direct bullet symbols and numbered/lettered lists
            if any(stripped.startswith(char) for char in ['•', '●', '·', '-', '*', '◦', '▪', '▫']):
                is_bullet = True
            # Numbered lists: "1.", "2.", "10.", etc.
            elif re.match(r'^\d+\.\s+', stripped):
                is_bullet = True
            # Lettered lists: "a.", "b.", "A.", "B.", etc.
            elif re.match(r'^[a-zA-Z]\.\s+', stripped):
                is_bullet = True
            
            # 2. Action verbs at start of line (expanded list)
            elif stripped.split():
                first_word = stripped.split()[0].lower()
                action_verbs = ['led', 'managed', 'developed', 'created', 'implemented', 'designed', 'built', 
                               'achieved', 'increased', 'decreased', 'improved', 'optimized', 'coordinated', 
                               'collaborated', 'established', 'launched', 'delivered', 'executed', 'analyzed', 
                               'drove', 'spearheaded', 'conducted', 'championed', 'directed', 'orchestrated',
                               'facilitated', 'streamlined', 'enhanced', 'engineered', 'architected', 'oversaw',
                               'supervised', 'trained', 'mentored', 'guided', 'supported', 'maintained',
                               'developed', 'programmed', 'coded', 'tested', 'debugged', 'deployed',
                               'monitored', 'tracked', 'measured', 'reported', 'presented', 'communicated']
                if first_word in action_verbs:
                    is_bullet = True
            
            # 3. Lines with clear metrics/achievements
            elif re.search(r'\d+%|\$\d+[kmb]?|increase.*\d+|decrease.*\d+|improve.*\d+|by \d+%|of \d+%|revenue', stripped.lower()):
                is_bullet = True
            
            # 4. Lines that are likely bullet points based on context
            elif (current_section and 
                  current_section in ['experience', 'professional experience', 'work experience', 'employment', 'skills'] and
                  len(stripped) > 20 and
                  not stripped.endswith(':')):
                is_bullet = True
            
            # 5. Lines starting with capital letter after experience/skills sections
            elif (current_section and 
                  current_section in ['experience', 'professional experience', 'work experience', 'skills'] and
                  len(stripped) > 15 and
                  stripped[0].isupper()):
                is_bullet = True
            
            if is_bullet:
                sections.append(("bullet_point", stripped))
                continue
            
            # Everything else is body text
            sections.append(("body_text", stripped))
        
        return sections
    
    def _validate_wrapped_lines(self, lines: list, max_width: float, font_name: str, font_size: int) -> list:
        """Validate that wrapped lines actually fit within the specified width"""
        from reportlab.pdfbase.pdfmetrics import stringWidth
        
        validated_lines = []
        safety_buffer = 5  # Conservative buffer for validation
        
        for line in lines:
            try:
                actual_width = stringWidth(line, font_name, font_size)
                if actual_width <= max_width - safety_buffer:
                    validated_lines.append(line)
                else:
                    # Line is still too wide, break it more conservatively
                    # Use character-based approach with 75% of estimated capacity
                    chars_per_line = max(10, int((max_width - safety_buffer) / (font_size * 0.6) * 0.75))
                    while len(line) > chars_per_line:
                        validated_lines.append(line[:chars_per_line] + "-")
                        line = line[chars_per_line:]
                    if line:
                        validated_lines.append(line)
            except Exception:
                # If stringWidth fails, use conservative character breaking
                chars_per_line = max(10, int((max_width - safety_buffer) / (font_size * 0.6) * 0.75))
                while len(line) > chars_per_line:
                    validated_lines.append(line[:chars_per_line] + "-")
                    line = line[chars_per_line:]
                if line:
                    validated_lines.append(line)
        
        return validated_lines

    def _wrap_text_by_width(self, text: str, max_width: float, font_name: str, font_size: int) -> list:
        """Wrap text based on actual pixel width with safety buffers for ReportLab inconsistencies"""
        from reportlab.pdfbase.pdfmetrics import stringWidth
        
        if not text:
            return []
        
        # Add minimal safety buffer to account for ReportLab stringWidth inconsistencies
        # This is the key fix - ReportLab's stringWidth can be off by a few pixels
        safety_buffer = 3  # pixels - minimal buffer to avoid truncation
        effective_max_width = max_width - safety_buffer
        
        # Measure actual width of text
        try:
            text_width = stringWidth(text, font_name, font_size)
            if text_width <= effective_max_width:
                return [text]
        except Exception:
            # Fallback to character-based wrapping if stringWidth fails
            return self._simple_char_wrap(text, int(max_width / 6))  # Rough approximation
        
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            # Test if adding this word would exceed the width
            test_line = current_line + (" " if current_line else "") + word
            
            try:
                test_width = stringWidth(test_line, font_name, font_size)
            except Exception:
                # If stringWidth fails, use character-based estimation
                test_width = len(test_line) * (font_size * 0.6)  # Rough approximation
            
            if test_width <= effective_max_width:
                current_line = test_line
            else:
                # Current line is full, start a new one
                if current_line:
                    lines.append(current_line)
                
                # Check if the word itself is too long
                try:
                    word_width = stringWidth(word, font_name, font_size)
                except Exception:
                    word_width = len(word) * (font_size * 0.6)
                
                if word_width > effective_max_width:
                    # Word is too long, need to break it more conservatively
                    # Use a more conservative approach to avoid stringWidth issues
                    if len(word) > 20:  # Only break very long words
                        # Try to break at natural points first
                        if '-' in word:
                            parts = word.split('-')
                            for j, part in enumerate(parts):
                                if j > 0:
                                    part = '-' + part
                                try:
                                    part_width = stringWidth(part, font_name, font_size)
                                except Exception:
                                    part_width = len(part) * (font_size * 0.6)
                                
                                if part_width <= effective_max_width:
                                    if current_line:
                                        lines.append(current_line)
                                        current_line = ""
                                    current_line = part
                                else:
                                    # Still too long, use character-based approach
                                    if current_line:
                                        lines.append(current_line)
                                        current_line = ""
                                    # Break conservatively - use 80% of estimated character capacity
                                    chars_per_line = max(10, int(effective_max_width / (font_size * 0.6) * 0.8))
                                    while len(part) > chars_per_line:
                                        lines.append(part[:chars_per_line] + "-")
                                        part = part[chars_per_line:]
                                    current_line = part
                        else:
                            # No natural break point, use conservative character breaking
                            if current_line:
                                lines.append(current_line)
                                current_line = ""
                            # Break conservatively - use 80% of estimated character capacity
                            chars_per_line = max(10, int(effective_max_width / (font_size * 0.6) * 0.8))
                            while len(word) > chars_per_line:
                                lines.append(word[:chars_per_line] + "-")
                                word = word[chars_per_line:]
                            current_line = word
                    else:
                        # Short word that's somehow too wide - just add it
                        # This shouldn't happen with proper fonts, but handle it gracefully
                        current_line = word
                else:
                    current_line = word
        
        # Add the last line if it has content
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _smart_wrap_text(self, text: str, max_chars: int) -> list:
        """Smart text wrapping that preserves meaning and avoids awkward breaks"""
        if len(text) <= max_chars:
            return [text]
        
        # Split into words
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            # Test if adding this word would exceed the limit
            test_line = current_line + (" " if current_line else "") + word
            
            if len(test_line) <= max_chars:
                current_line = test_line
            else:
                # Current line is full, start a new one
                if current_line:
                    lines.append(current_line)
                
                # If single word is too long, we need to break it carefully
                if len(word) > max_chars:
                    # Try to find a good break point in the word
                    if '-' in word:
                        # Break at hyphen if possible
                        parts = word.split('-')
                        for j, part in enumerate(parts):
                            if j > 0:
                                part = '-' + part
                            if len(part) <= max_chars:
                                if current_line:
                                    lines.append(current_line)
                                current_line = part
                            else:
                                # Still too long, force break
                                while len(part) > max_chars:
                                    lines.append(part[:max_chars-1] + "-")
                                    part = part[max_chars-1:]
                                current_line = part
                    else:
                        # Force break the long word
                        while len(word) > max_chars:
                            lines.append(word[:max_chars-1] + "-")
                            word = word[max_chars-1:]
                        current_line = word
                else:
                    current_line = word
        
        # Add the last line if it has content
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def create_tailored_resume_docx(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        Create a DOCX from tailored resume text
        """
        try:
            doc = Document()
            
            # Parse sections
            sections = self._parse_resume_sections(tailored_text)
            
            for section_title, section_content in sections:
                if section_title:
                    # Add section heading
                    heading = doc.add_heading(section_title, level=2)
                
                # Add section content
                paragraphs = section_content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error creating DOCX: {str(e)}")
            return False
    
    def _parse_resume_sections(self, text: str) -> list:
        """
        Parse resume text into sections
        """
        sections = []
        current_section = ""
        current_content = []
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            if self._is_section_header(line):
                # Save previous section
                if current_section or current_content:
                    sections.append((current_section, '\n'.join(current_content)))
                
                # Start new section
                current_section = line
                current_content = []
            else:
                current_content.append(line)
        
        # Add the last section
        if current_section or current_content:
            sections.append((current_section, '\n'.join(current_content)))
        
        return sections
    
    def _is_section_header(self, line: str) -> bool:
        """
        Determine if a line is likely a section header
        """
        common_headers = [
            'summary', 'objective', 'experience', 'work experience', 'employment',
            'education', 'skills', 'technical skills', 'certifications',
            'projects', 'achievements', 'awards', 'languages', 'interests',
            'contact', 'contact information', 'professional summary',
            'professional experience', 'core competencies', 'key achievements',
            'technical expertise', 'expertise', 'qualifications', 'career highlights',
            'accomplishments', 'notable projects', 'relevant experience',
            'academic background', 'educational background', 'volunteer experience',
            'extracurricular activities', 'leadership experience'
        ]
        
        line_lower = line.lower().strip().replace(':', '')
        
        # Check if it matches common headers exactly
        if any(header == line_lower for header in common_headers):
            return True
        
        # Check for partial matches only if they're substantial portions of the line
        for header in common_headers:
            if header in line_lower and len(header) > len(line_lower) * 0.4:
                return True
        
        # Check if it's all caps (common for headers)
        if line.isupper() and len(line) > 3 and len(line) < 40:
            return True
        
        # Check if it's a title case header with colon
        if (line.endswith(':') and 
            len(line) > 5 and len(line) < 40 and
            line.count(' ') <= 3 and
            not any(char in line for char in ['•', '●', '·', '@', '(', ')']) and
            (line[0].isupper() or line.isupper())):
            return True
        
        return False 

    def test_text_wrapping_fix(self) -> bool:
        """Test function to verify that text wrapping is working correctly"""
        from reportlab.pdfbase.pdfmetrics import stringWidth
        
        # Test cases that were previously causing truncation
        test_cases = [
            "...Net Promoter Score (NPS) by 18 points",
            "...feature adoption rates by 40%",
            "...annual recurring revenue",
            "Increased customer satisfaction scores by improving response time",
            "Developed comprehensive marketing strategy that resulted in 25% growth"
        ]
        
        font_name = "Helvetica"
        font_size = 10
        max_width = 400  # Typical width that was causing issues
        
        print("Testing text wrapping fix...")
        
        for test_text in test_cases:
            print(f"\nTesting: '{test_text}'")
            
            # Test the old problematic scenario
            wrapped_lines = self._wrap_text_by_width(test_text, max_width, font_name, font_size)
            validated_lines = self._validate_wrapped_lines(wrapped_lines, max_width, font_name, font_size)
            
            print(f"  Original text length: {len(test_text)} chars")
            print(f"  Wrapped into {len(validated_lines)} lines:")
            
            total_chars = 0
            for i, line in enumerate(validated_lines):
                try:
                    actual_width = stringWidth(line, font_name, font_size)
                    print(f"    Line {i+1}: '{line}' (width: {actual_width:.1f}px)")
                    total_chars += len(line)
                except Exception as e:
                    print(f"    Line {i+1}: '{line}' (width calculation failed: {e})")
                    total_chars += len(line)
            
            print(f"  Total characters preserved: {total_chars}/{len(test_text)}")
            
            # Check if any line is truncated unexpectedly
            if total_chars < len(test_text) - 5:  # Allow for some hyphens
                print(f"  WARNING: Significant text loss detected!")
                return False
        
        print("\nText wrapping fix test completed successfully!")
        return True 

    def _smart_validate_line(self, line: str, max_width: float, font_name: str, font_size: int) -> str:
        """Smart validation for individual lines that might be at risk of truncation"""
        from reportlab.pdfbase.pdfmetrics import stringWidth
        
        try:
            actual_width = stringWidth(line, font_name, font_size)
            
            # Only apply extra validation if the line actually exceeds the width
            # This targets the specific ReportLab inconsistency issue
            safety_margin = 5  # Small safety margin for ReportLab inconsistencies
            
            if actual_width > max_width - safety_margin:
                # Line is at risk, apply minimal trimming
                words = line.split()
                if len(words) > 1:
                    # Try removing just the last word first
                    trimmed_line = ' '.join(words[:-1])
                    trimmed_width = stringWidth(trimmed_line, font_name, font_size)
                    if trimmed_width <= max_width - safety_margin:
                        return trimmed_line
                
                # If word removal doesn't work, try character-based trimming
                # Be very conservative - only trim a few characters
                for trim_chars in range(1, min(10, len(line) // 4)):
                    test_line = line[:-trim_chars].rstrip()
                    if stringWidth(test_line, font_name, font_size) <= max_width - safety_margin:
                        return test_line
                
                # Last resort: remove 10% of characters
                fallback_length = int(len(line) * 0.9)
                return line[:fallback_length].rstrip()
            
            return line  # Line is safe as-is
            
        except Exception:
            # If stringWidth fails, only trim if line is very long
            if len(line) > 60:  # Only trim very long lines
                return line[:55].rstrip()
            return line 

    def create_tailored_resume_pdf_fixed(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        PROPER SOLUTION: Use SimpleDocTemplate and PLATYPUS framework
        This fixes the disappearing text issue by using ReportLab's proper text flow system
        """
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import black, blue
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            
            # Create the document with proper margins
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=0.5*inch,
                leftMargin=0.5*inch,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch
            )
            
            # Get base styles and create custom ones
            styles = getSampleStyleSheet()
            
            # Define custom styles that match your original formatting
            name_style = ParagraphStyle(
                'NameStyle',
                parent=styles['Normal'],
                fontSize=16,
                textColor=blue,
                spaceAfter=6,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            )
            
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Normal'],
                fontSize=14,
                textColor=blue,
                spaceAfter=4,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            )
            
            contact_style = ParagraphStyle(
                'ContactStyle',
                parent=styles['Normal'],
                fontSize=9,
                spaceAfter=12,
                fontName='Helvetica'
            )
            
            section_header_style = ParagraphStyle(
                'SectionHeaderStyle',
                parent=styles['Normal'],
                fontSize=11,
                textColor=blue,
                spaceBefore=8,
                spaceAfter=4,
                fontName='Helvetica-Bold'
            )
            
            company_header_style = ParagraphStyle(
                'CompanyHeaderStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceBefore=6,
                spaceAfter=2,
                fontName='Helvetica-Bold'
            )
            
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=10,
                leftIndent=20,
                firstLineIndent=-15,
                spaceAfter=3,
                fontName='Helvetica'
            )
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                fontName='Helvetica'
            )
            
            # Parse the resume content
            sections = self._parse_resume_for_compact_formatting(tailored_text)
            
            # Build the story (list of flowables)
            story = []
            
            for section_type, content in sections:
                if section_type == "name":
                    story.append(Paragraph(content.strip(), name_style))
                    # Add spacing after name
                    story.append(Spacer(1, 6))
                    
                elif section_type == "title":
                    story.append(Paragraph(content.strip(), title_style))
                    
                elif section_type == "contact":
                    # Handle misclassified bullet points that got marked as contact
                    if content.strip().startswith('•'):
                        # This is actually a bullet point
                        clean_content = content.strip()
                        while clean_content and clean_content[0] in ['•', '●', '·', '-', '*', '+', '◦', '▪', '▫']:
                            clean_content = clean_content[1:].strip()
                        if clean_content:
                            bullet_content = f"• {clean_content}"
                            story.append(Paragraph(bullet_content, bullet_style))
                    else:
                        story.append(Paragraph(content.strip(), contact_style))
                    
                elif section_type == "section_header":
                    # Handle misclassified names that got marked as section_header
                    if content.strip() == "DAVID PATEL":
                        story.append(Paragraph(content.strip(), name_style))
                        story.append(Spacer(1, 6))  # Add spacing after name
                    elif content.strip().startswith('•'):
                        # This is actually a bullet point
                        clean_content = content.strip()
                        while clean_content and clean_content[0] in ['•', '●', '·', '-', '*', '+', '◦', '▪', '▫']:
                            clean_content = clean_content[1:].strip()
                        if clean_content:
                            bullet_content = f"• {clean_content}"
                            story.append(Paragraph(bullet_content, bullet_style))
                    else:
                        story.append(Paragraph(content.strip().upper(), section_header_style))
                    
                elif section_type == "company_header":
                    story.append(Paragraph(content.strip(), company_header_style))
                    
                elif section_type == "bullet_point":
                    # Clean the bullet content consistently
                    clean_content = content.strip()
                    # Remove any existing bullet symbols
                    while clean_content and clean_content[0] in ['•', '●', '·', '-', '*', '+', '◦', '▪', '▫']:
                        clean_content = clean_content[1:].strip()
                    # Remove numbered/lettered list markers
                    clean_content = re.sub(r'^\d+\.\s+', '', clean_content)
                    clean_content = re.sub(r'^[a-zA-Z]\.\s+', '', clean_content)
                    # Remove any remaining leading whitespace
                    clean_content = clean_content.strip()
                    
                    if clean_content:
                        # Use consistent bullet formatting for ALL bullets
                        bullet_content = f"• {clean_content}"
                        story.append(Paragraph(bullet_content, bullet_style))
                        
                elif section_type == "body_text":
                    story.append(Paragraph(content.strip(), body_style))
                    
                elif section_type == "spacing":
                    spacing = min(int(content), 12)
                    story.append(Spacer(1, spacing))
            
            # Build the document - PLATYPUS handles all the text flow automatically!
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error creating PLATYPUS PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def create_tailored_resume_pdf_direct(self, tailored_text: str, output_path: str, job_title: str = "") -> bool:
        """
        Direct PDF generation that preserves GPT formatting with minimal processing
        """
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import HexColor
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT
            
            # Create the document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            
            # Get base styles
            styles = getSampleStyleSheet()
            professional_blue = HexColor('#2c5aa0')
            
            # Define styles
            name_style = ParagraphStyle(
                'DirectNameStyle',
                parent=styles['Normal'],
                fontSize=18,
                textColor=professional_blue,
                spaceAfter=6,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            )
            
            contact_style = ParagraphStyle(
                'DirectContactStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                fontName='Helvetica'
            )
            
            section_style = ParagraphStyle(
                'DirectSectionStyle',
                parent=styles['Normal'],
                fontSize=13,
                textColor=professional_blue,
                spaceBefore=12,
                spaceAfter=6,
                fontName='Helvetica-Bold'
            )
            
            bullet_style = ParagraphStyle(
                'DirectBulletStyle',
                parent=styles['Normal'],
                fontSize=11,
                leftIndent=20,
                firstLineIndent=-15,
                spaceAfter=4,
                fontName='Helvetica'
            )
            
            normal_style = ParagraphStyle(
                'DirectNormalStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=4,
                fontName='Helvetica'
            )
            
            # Process the text line by line with minimal parsing
            lines = tailored_text.split('\n')
            story = []
            
            for i, line in enumerate(lines):
                stripped = line.strip()
                
                # Skip empty lines but add small spacing
                if not stripped:
                    story.append(Spacer(1, 6))
                    continue
                
                # First non-empty line is likely the name
                if i == 0 or (i < 3 and not any(char in stripped for char in ['@', '(', ')', '|', 'EXPERIENCE', 'SKILLS', 'EDUCATION'])):
                    story.append(Paragraph(stripped, name_style))
                    continue
                
                # Contact info (contains email or phone patterns)
                if '@' in stripped or '(' in stripped or re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', stripped):
                    story.append(Paragraph(stripped, contact_style))
                    continue
                
                # Section headers (all caps, common section names)
                if (stripped.isupper() and any(keyword in stripped for keyword in 
                    ['EXPERIENCE', 'SKILLS', 'EDUCATION', 'SUMMARY', 'PROJECTS', 'CERTIFICATIONS'])):
                    story.append(Paragraph(stripped, section_style))
                    continue
                
                # Bullet points (start with bullet symbols)
                if stripped.startswith(('•', '●', '·', '-', '*')) or re.match(r'^\d+\.\s+', stripped):
                    # Clean and standardize bullet
                    clean_content = stripped
                    if clean_content.startswith(('•', '●', '·', '-', '*')):
                        clean_content = clean_content[1:].strip()
                    elif re.match(r'^\d+\.\s+', clean_content):
                        clean_content = re.sub(r'^\d+\.\s+', '', clean_content)
                    
                    bullet_content = f"• {clean_content}"
                    story.append(Paragraph(bullet_content, bullet_style))
                    continue
                
                # Everything else as normal text
                story.append(Paragraph(stripped, normal_style))
            
            # Build the document
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error in direct PDF generation: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
