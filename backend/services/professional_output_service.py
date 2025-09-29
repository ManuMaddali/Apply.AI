"""
Professional Output Service - Phase 2 Advanced Features
Provides ATS-optimized PDF/Word generation with professional templates
"""

from typing import Dict, List, Optional, Tuple, Any
import os
import re
from io import BytesIO
from datetime import datetime
import asyncio


def _prefer_backend(module_path: str, fallback_path: str, attr: str | None = None):
    try:
        mod = __import__(module_path, fromlist=['*'])
    except ImportError:
        mod = __import__(fallback_path, fromlist=['*'])
    return getattr(mod, attr) if attr else mod

# Import existing components
try:
    from utils.resume_editor import ResumeEditor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.colors import HexColor, black, white
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DEPENDENCIES_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Some dependencies not available: {e}")
    DEPENDENCIES_AVAILABLE = False

# New template engine (HTML/CSS -> PDF)
TemplateEngine = _prefer_backend('backend.services.template_engine', 'services.template_engine', attr='TemplateEngine')
parse_resume_text_to_schema = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='parse_resume_text_to_schema')
Resume = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='Resume')
Contact = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='Contact')
ExperienceItem = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='ExperienceItem')
EducationItem = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='EducationItem')
Skill = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='Skill')
ProjectItem = _prefer_backend('backend.models.resume_schema', 'models.resume_schema', attr='ProjectItem')
clean_and_compact = _prefer_backend('backend.services.cleaners', 'services.cleaners', attr='clean_and_compact')

# Import our comprehensive resume parser
try:
    ResumeParser = _prefer_backend('backend.services.resume_parser', 'services.resume_parser', attr='ResumeParser')
    RESUME_PARSER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Resume parser not available: {e}")
    RESUME_PARSER_AVAILABLE = False

# Keep imports of AdvancedFormattingService available for backward compatibility,
# but they are deprecated and not used for PDF generation anymore.
ADVANCED_FMT_AVAILABLE = True
try:
    advanced_module = _prefer_backend('backend.services.advanced_formatting_service', 'services.advanced_formatting_service')
    AdvancedFormattingService = advanced_module.AdvancedFormattingService
    FormattingOptions = advanced_module.FormattingOptions
    FormattingTemplate = advanced_module.FormattingTemplate
    ColorScheme = advanced_module.ColorScheme
    FontFamily = advanced_module.FontFamily
except Exception as e:
    print(f"⚠️ AdvancedFormattingService not available: {e}")
    ADVANCED_FMT_AVAILABLE = False

class ATSScorer:
    """ATS Compatibility Scoring System"""
    
    def __init__(self):
        self.ats_keywords = {
            'technical': ['python', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 'kubernetes'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem-solving', 'analytical'],
            'experience': ['years', 'experience', 'developed', 'managed', 'led', 'implemented'],
            'education': ['degree', 'bachelor', 'master', 'university', 'college', 'certification']
        }
    
    def calculate_ats_score(self, resume_text: str, job_description: str = "") -> Dict[str, Any]:
        """Calculate comprehensive ATS compatibility score (0-100)"""
        try:
            score_components = {
                'keyword_match': self._score_keyword_match(resume_text, job_description),
                'formatting': self._score_formatting(resume_text),
                'structure': self._score_structure(resume_text),
                'readability': self._score_readability(resume_text),
                'completeness': self._score_completeness(resume_text)
            }
            
            # Weighted scoring
            weights = {
                'keyword_match': 0.35,
                'formatting': 0.20,
                'structure': 0.20,
                'readability': 0.15,
                'completeness': 0.10
            }
            
            total_score = sum(score_components[key] * weights[key] for key in weights)
            
            return {
                'total_score': min(100, max(0, int(total_score))),
                'components': score_components,
                'recommendations': self._generate_recommendations(score_components),
                'grade': self._get_grade(total_score)
            }
        except Exception as e:
            return {'total_score': 75, 'error': str(e), 'grade': 'B'}
    
    def _score_keyword_match(self, resume_text: str, job_description: str) -> float:
        """Score keyword matching with job description"""
        if not job_description:
            return 70.0  # Default score when no job description
        
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        # Extract keywords from job description
        job_keywords = set(re.findall(r'\b\w+\b', job_lower))
        job_keywords = {word for word in job_keywords if len(word) > 3}
        
        # Count matches
        matches = sum(1 for keyword in job_keywords if keyword in resume_lower)
        match_ratio = matches / len(job_keywords) if job_keywords else 0
        
        return min(100, match_ratio * 100)
    
    def _score_formatting(self, resume_text: str) -> float:
        """Score formatting quality"""
        score = 100.0
        
        # Check for common formatting issues
        if '\\bullet' in resume_text:
            score -= 15  # Literal bullet text
        if resume_text.count('\n\n') < 3:
            score -= 10  # Insufficient spacing
        if len(re.findall(r'[A-Z]{2,}', resume_text)) > 10:
            score -= 5   # Too much ALL CAPS
        
        return max(0, score)
    
    def _score_structure(self, resume_text: str) -> float:
        """Score resume structure"""
        score = 0.0
        sections = ['experience', 'education', 'skills', 'summary']
        
        for section in sections:
            if any(keyword in resume_text.lower() for keyword in [section, section.replace('e', 'a')]):
                score += 25
        
        return min(100, score)
    
    def _score_readability(self, resume_text: str) -> float:
        """Score readability"""
        lines = resume_text.split('\n')
        avg_line_length = sum(len(line) for line in lines) / len(lines) if lines else 0
        
        # Optimal line length is 60-80 characters
        if 60 <= avg_line_length <= 80:
            return 100.0
        elif 40 <= avg_line_length <= 100:
            return 80.0
        else:
            return 60.0
    
    def _score_completeness(self, resume_text: str) -> float:
        """Score resume completeness"""
        required_elements = ['email', 'phone', 'experience', 'skills']
        score = 0.0
        
        for element in required_elements:
            if element in resume_text.lower():
                score += 25
        
        return min(100, score)
    
    def _generate_recommendations(self, scores: Dict[str, float]) -> List[str]:
        """Generate improvement recommendations"""
        recommendations = []
        
        if scores['keyword_match'] < 70:
            recommendations.append("Add more job-specific keywords from the job description")
        if scores['formatting'] < 80:
            recommendations.append("Improve formatting consistency and remove literal bullet text")
        if scores['structure'] < 80:
            recommendations.append("Ensure all major sections (Experience, Education, Skills) are present")
        if scores['readability'] < 70:
            recommendations.append("Optimize line length and paragraph structure for better readability")
        
        return recommendations
    
    def _get_grade(self, score: float) -> str:
        """Convert numeric score to letter grade"""
        if score >= 90: return 'A+'
        elif score >= 85: return 'A'
        elif score >= 80: return 'B+'
        elif score >= 75: return 'B'
        elif score >= 70: return 'C+'
        elif score >= 65: return 'C'
        else: return 'D'

class ProfessionalTemplateEngine:
    """Professional template system with multiple ATS-optimized layouts"""
    
    def __init__(self):
        self.templates = {
            'executive': self._executive_template,
            'technical': self._technical_template,
            'creative': self._creative_template,
            'modern': self._modern_template,
            'classic': self._classic_template
        }
    
    def get_template_styles(self, template_name: str) -> Dict[str, ParagraphStyle]:
        """Get styles for specific template"""
        if template_name not in self.templates:
            template_name = 'modern'
        
        return self.templates[template_name]()
    
    def _executive_template(self) -> Dict[str, ParagraphStyle]:
        """Executive template - Conservative, professional"""
        blue = HexColor('#1f4e79')
        return {
            'name': ParagraphStyle('Name', fontName='Helvetica-Bold', fontSize=20, textColor=blue, spaceAfter=3),
            'title': ParagraphStyle('Title', fontName='Helvetica', fontSize=14, textColor=blue, spaceAfter=8),
            'contact': ParagraphStyle('Contact', fontName='Helvetica', fontSize=10, spaceAfter=12),
            'section': ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=12, textColor=blue, spaceBefore=15, spaceAfter=8),
            'body': ParagraphStyle('Body', fontName='Helvetica', fontSize=10, spaceAfter=6),
            'bullet': ParagraphStyle('Bullet', fontName='Helvetica', fontSize=10, leftIndent=15, firstLineIndent=-15, spaceAfter=4)
        }
    
    def _technical_template(self) -> Dict[str, ParagraphStyle]:
        """Technical template - Clean, modern"""
        green = HexColor('#2d5016')
        return {
            'name': ParagraphStyle('Name', fontName='Helvetica-Bold', fontSize=18, textColor=green, spaceAfter=2),
            'title': ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=13, textColor=green, spaceAfter=6),
            'contact': ParagraphStyle('Contact', fontName='Helvetica', fontSize=9, spaceAfter=10),
            'section': ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=11, textColor=green, spaceBefore=12, spaceAfter=6),
            'body': ParagraphStyle('Body', fontName='Helvetica', fontSize=10, spaceAfter=5),
            'bullet': ParagraphStyle('Bullet', fontName='Helvetica', fontSize=10, leftIndent=12, firstLineIndent=-12, spaceAfter=3)
        }
    
    def _creative_template(self) -> Dict[str, ParagraphStyle]:
        """Creative template - Modern with color"""
        purple = HexColor('#663399')
        return {
            'name': ParagraphStyle('Name', fontName='Helvetica-Bold', fontSize=19, textColor=purple, spaceAfter=3),
            'title': ParagraphStyle('Title', fontName='Helvetica-Oblique', fontSize=13, textColor=purple, spaceAfter=7),
            'contact': ParagraphStyle('Contact', fontName='Helvetica', fontSize=10, spaceAfter=11),
            'section': ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=12, textColor=purple, spaceBefore=14, spaceAfter=7),
            'body': ParagraphStyle('Body', fontName='Helvetica', fontSize=10, spaceAfter=6),
            'bullet': ParagraphStyle('Bullet', fontName='Helvetica', fontSize=10, leftIndent=14, firstLineIndent=-14, spaceAfter=4)
        }
    
    def _modern_template(self) -> Dict[str, ParagraphStyle]:
        """Modern template - Your current reference style"""
        blue = HexColor('#4472C4')
        return {
            'name': ParagraphStyle('Name', fontName='Helvetica-Bold', fontSize=18, textColor=blue, spaceAfter=2),
            'title': ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=14, textColor=blue, spaceAfter=6),
            'contact': ParagraphStyle('Contact', fontName='Helvetica', fontSize=10, spaceAfter=12),
            'section': ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=11, textColor=blue, spaceBefore=12, spaceAfter=6),
            'body': ParagraphStyle('Body', fontName='Helvetica', fontSize=10, spaceAfter=6),
            'bullet': ParagraphStyle('Bullet', fontName='Helvetica', fontSize=10, leftIndent=12, firstLineIndent=-12, spaceAfter=3)
        }
    
    def _classic_template(self) -> Dict[str, ParagraphStyle]:
        """Classic template - Traditional black and white"""
        return {
            'name': ParagraphStyle('Name', fontName='Times-Bold', fontSize=18, spaceAfter=3),
            'title': ParagraphStyle('Title', fontName='Times-Italic', fontSize=14, spaceAfter=8),
            'contact': ParagraphStyle('Contact', fontName='Times-Roman', fontSize=10, spaceAfter=12),
            'section': ParagraphStyle('Section', fontName='Times-Bold', fontSize=12, spaceBefore=15, spaceAfter=8),
            'body': ParagraphStyle('Body', fontName='Times-Roman', fontSize=10, spaceAfter=6),
            'bullet': ParagraphStyle('Bullet', fontName='Times-Roman', fontSize=10, leftIndent=15, firstLineIndent=-15, spaceAfter=4)
        }

class ProfessionalOutputService:
    """Main service for professional PDF/Word output with ATS optimization"""
    
    def __init__(self):
        self.resume_editor = ResumeEditor() if DEPENDENCIES_AVAILABLE else None
        self.ats_scorer = ATSScorer()
        self.template_engine = ProfessionalTemplateEngine()
    
    def _convert_parsed_to_resume_schema(self, parsed_data: Dict[str, Any]) -> Resume:
        """Convert parser output to Resume schema format."""
        try:
            # Extract contact information
            contact = Contact(
                email=parsed_data.get('email', ''),
                phone=parsed_data.get('phone', ''),
                location=parsed_data.get('location', ''),
                links=[]
            )
            
            # Add LinkedIn and GitHub as links if available
            if parsed_data.get('linkedin'):
                contact.links.append({'label': 'LinkedIn', 'url': parsed_data['linkedin']})
            if parsed_data.get('github'):
                contact.links.append({'label': 'GitHub', 'url': parsed_data['github']})
            
            # Convert experience
            experience_items = []
            for exp in parsed_data.get('experience', []):
                bullets = exp.get('responsibilities', [])
                # Ensure bullets are strings
                if isinstance(bullets, list):
                    bullets = [str(b) for b in bullets if b]
                
                experience_items.append(ExperienceItem(
                    title=exp.get('title', ''),
                    company=exp.get('company', ''),
                    dates=exp.get('dates', ''),
                    bullets=bullets
                ))
            
            # Convert education
            education_items = []
            for edu in parsed_data.get('education', []):
                school = edu.get('institution', '')
                degree = edu.get('degree', '')
                field = edu.get('field', '')
                year = edu.get('year', '')
                
                # Combine into a school string
                school_str = f"{degree} {field}" if degree and field else (degree or field or '')
                if school:
                    school_str = f"{school} - {school_str}" if school_str else school
                if year:
                    school_str = f"{school_str} ({year})" if school_str else year
                
                if school_str:
                    education_items.append(EducationItem(school=school_str))
            
            # Convert skills
            skills_list = []
            skills_dict = parsed_data.get('skills', {})
            if isinstance(skills_dict, dict):
                # Combine all skill categories
                for category in ['technical', 'languages', 'tools', 'soft', 'other']:
                    skills_list.extend(skills_dict.get(category, []))
            elif isinstance(skills_dict, list):
                skills_list = skills_dict
            
            # Remove duplicates while preserving order
            seen = set()
            unique_skills = []
            for skill in skills_list:
                if skill and skill.lower() not in seen:
                    seen.add(skill.lower())
                    unique_skills.append(Skill(name=skill))
            
            # Convert projects
            project_items = []
            for proj in parsed_data.get('projects', []):
                name = proj.get('name', 'Project')
                desc = proj.get('description', '')
                bullets = [desc] if desc else []
                project_items.append(ProjectItem(name=name, bullets=bullets))
            
            # Get headline from summary (first sentence or first 80 chars)
            summary = parsed_data.get('summary', '')
            headline = None
            if summary:
                # Try to get first sentence
                first_sentence = summary.split('.')[0]
                if len(first_sentence) <= 100:
                    headline = first_sentence
                else:
                    headline = summary[:80] + '...'
            
            return Resume(
                name=parsed_data.get('name', ''),
                headline=headline,
                contact=contact,
                summary=summary,
                experience=experience_items,
                projects=project_items,
                education=education_items,
                skills=unique_skills
            )
            
        except Exception as e:
            print(f"⚠️ Error converting parsed data to Resume schema: {e}")
            # Return empty Resume on error
            return Resume()
    
    def generate_professional_pdf(
        self,
        resume_text: str,
        job_description: str = "",
        template: str = "executive_compact",
        ats_optimize: bool = True
    ) -> Dict[str, Any]:
        """Generate professional PDF using HTML/CSS templates rendered via Chromium.
        
        Uses the comprehensive resume parser to extract structured data,
        then renders clean PDFs without overlapping text.
        """
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            logger.info(f"🚀 Starting professional PDF generation with template: {template}")
            
            # Calculate ATS score (kept for compatibility)
            ats_results = self.ats_scorer.calculate_ats_score(resume_text, job_description)
            logger.info(f"📊 ATS Score calculated: {ats_results.get('total_score', 0)}")

            # Optional text optimizations
            optimized_text = self._apply_ats_optimizations(resume_text, ats_results) if ats_optimize else resume_text

            # Parse resume with comprehensive parser if available
            if RESUME_PARSER_AVAILABLE:
                logger.info("📝 Using comprehensive resume parser")
                parser = ResumeParser()
                # Sanitize input of JD-derived junk before parsing
                try:
                    sanitize_input_text = _prefer_backend('services.content_filters', 'backend.services.content_filters', attr='sanitize_input_text')
                except Exception:
                    sanitize_input_text = lambda x: x
                parsed_data = parser.parse(sanitize_input_text(optimized_text))
                
                # Validate parsed data
                validation_errors = parser.validate_parsed_data(parsed_data)
                if validation_errors:
                    logger.warning(f"⚠️ Validation issues: {validation_errors}")
                
                # Convert to Resume schema
                resume_obj = self._convert_parsed_to_resume_schema(parsed_data)
                if not resume_obj or not resume_obj.name:
                    logger.warning("⚠️ Parsed resume missing critical fields; falling back to raw text output")
                else:
                    logger.info(f"✅ Successfully parsed resume: {resume_obj.name}")
            else:
                logger.info("📝 Using fallback resume parser")
                # Fallback to original parser
                try:
                    sanitize_input_text = _prefer_backend('services.content_filters', 'backend.services.content_filters', attr='sanitize_input_text')
                except Exception:
                    sanitize_input_text = lambda x: x
                resume_obj = parse_resume_text_to_schema(sanitize_input_text(optimized_text))
            
            resume_json = clean_and_compact(resume_obj.dict())

            # Generate formatted display text using parsed data when available
            formatted_display_text = make_short_preview_string(resume_json)
            
            # Add detailed formatting for better display
            if RESUME_PARSER_AVAILABLE and resume_obj.name:
                # Create a more detailed display text
                display_parts = []
                display_parts.append(f"👤 {resume_obj.name}")
                
                if resume_obj.contact.email:
                    display_parts.append(f"📧 {resume_obj.contact.email}")
                if resume_obj.contact.phone:
                    display_parts.append(f"📱 {resume_obj.contact.phone}")
                if resume_obj.contact.location:
                    display_parts.append(f"📍 {resume_obj.contact.location}")
                
                if resume_obj.headline:
                    display_parts.append(f"\n💼 {resume_obj.headline}")
                
                if resume_obj.experience:
                    display_parts.append(f"\n🏢 Experience: {len(resume_obj.experience)} positions")
                    for exp in resume_obj.experience[:2]:  # Show first 2 experiences
                        if exp.title and exp.company:
                            display_parts.append(f"  • {exp.title} at {exp.company}")
                
                if resume_obj.education:
                    display_parts.append(f"\n🎓 Education: {len(resume_obj.education)} entries")
                
                if resume_obj.skills:
                    display_parts.append(f"\n🔧 Skills: {len(resume_obj.skills)} items")
                    skill_names = [s.name for s in resume_obj.skills[:5]]  # First 5 skills
                    if skill_names:
                        display_parts.append(f"  {', '.join(skill_names)}...")
                
                formatted_display_text = "\n".join(display_parts)

            # Produce PDF using template engine pipeline
            logger.info("📄 Generating PDF from TemplateEngine")
            bundle_name = template or "executive_compact"
            pdf_bytes = b""
            try:
                pdf_bytes = TemplateEngine.render_pdf_sync(
                    template_id=bundle_name,
                    resume_json=resume_json,
                    resume_text=optimized_text,
                    bundle=bundle_name,
                )
            except Exception as render_err:
                logger.error(f"❌ TemplateEngine rendering failed, falling back to direct ReportLab: {render_err}", exc_info=True)
                try:
                    reportlab_direct = _prefer_backend(
                        'backend.services.reportlab_direct',
                        'services.reportlab_direct'
                    )
                    fallback_text = resume_text or formatted_display_text or optimized_text
                    pdf_bytes = reportlab_direct.generate_pdf_from_text(
                        fallback_text,
                        template=bundle_name,
                        display_text=formatted_display_text,
                    )
                    logger.info("✅ Fallback ReportLab PDF generated successfully")
                except Exception as direct_err:
                    logger.error(f"❌ ReportLab fallback failed: {direct_err}", exc_info=True)
                    raise render_err
            
            if not pdf_bytes:
                raise ValueError("PDF generation returned empty content")
            
            logger.info(f"✅ PDF generated successfully: {len(pdf_bytes)} bytes")

            # Keep ATS scoring out of resume content; return separately so callers can persist as scorecard
            response_payload = {
                'success': True,
                'pdf_content': pdf_bytes,
                'formatted_text': formatted_display_text,
                'template_used': template or "executive_compact",
                'optimizations_applied': ats_optimize,
                'parser_used': 'comprehensive' if RESUME_PARSER_AVAILABLE else 'fallback'
            }

            return response_payload, ats_results
            
        except Exception as e:
            logger.error(f"❌ Professional PDF generation failed: {e}", exc_info=True)
            print(f"❌ Professional PDF generation failed: {e}")

            # Return error with details
            return {
                'success': False,
                'error': str(e),
                'error_type': type(e).__name__,
                'template_used': template or "executive_compact",
                'parser_used': 'comprehensive' if RESUME_PARSER_AVAILABLE else 'fallback'
            }, None
    
    def _generate_formatted_display_text(self, resume_text: str, template: str) -> str:
        """Generate properly formatted display text that matches template styling"""
        try:
            print(f"🎨 Starting display formatting for template: {template}")
            print(f"📝 Resume text length: {len(resume_text)}")
            print(f"📝 Resume preview (first 200 chars): {resume_text[:200]}...")
            
            # Parse resume into structured sections
            sections = self._parse_resume_sections(resume_text)
            print(f"🔍 Sections parsed: {list(sections.keys())}")
            
            # Apply template-specific formatting
            if template.lower() == "modern":
                formatted = self._format_modern_template(sections)
            elif template.lower() == "classic":
                formatted = self._format_classic_template(sections)
            elif template.lower() == "technical":
                formatted = self._format_technical_template(sections)
            else:
                formatted = self._format_modern_template(sections)  # Default to modern
            
            print(f"✅ Display formatting successful! Length: {len(formatted)}")
            print(f"🎨 Template preview (first 300 chars): {formatted[:300]}...")
            return formatted
                
        except Exception as e:
            print(f"❌ Display formatting failed: {e}")
            import traceback
            print(f"🔍 Full error traceback: {traceback.format_exc()}")
            return resume_text  # Fallback to original
    
    def _parse_resume_sections(self, resume_text: str) -> Dict[str, Any]:
        """
        Parse resume text into structured sections:
        - Extracts name, contact info, and main sections (experience, education, skills, summary, other)
        - Handles common section header variations and preserves formatting
        """
        import re

        print(f"🔍 _parse_resume_sections called with text length: {len(resume_text)}")
        print(f"🔍 Resume text preview: {resume_text[:150]}...")

        # Section header patterns (case-insensitive)
        SECTION_HEADERS = {
            "experience": re.compile(r"^(work\s+)?experience|professional\s+experience|employment\s+history$", re.I),
            "education": re.compile(r"^education|academic\s+background$", re.I),
            "skills": re.compile(r"^skills|technical\s+skills|core\s+competencies$", re.I),
            "summary": re.compile(r"^summary|professional\s+summary|profile|about\s+me$", re.I),
        }
        # For "other" section, anything not matching above

        # Contact info patterns
        EMAIL_PATTERN = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
        PHONE_PATTERN = re.compile(r"\b(\+?\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?){1,2}\d{3,4}\b")
        LOCATION_KEYWORDS = [
            "street", "st.", "ave", "avenue", "road", "rd.", "blvd", "suite", "floor",
            "city", "state", "zip", "country", "usa", "united states", "canada",
            "atlanta", "new york", "san francisco", "los angeles", "chicago", "boston", "ga", "ny", "ca", "il", "ma"
        ]

        try:
            lines = [line.strip() for line in resume_text.strip().split('\n') if line.strip()]
            print(f"🔍 Parsed {len(lines)} lines from resume")

            if not lines:
                print("⚠️ Empty resume text for parsing")
                return {
                    'header': 'Professional Resume',
                    'experience': '',
                    'education': '',
                    'skills': '',
                    'summary': '',
                    'contact': '',
                    'other': ''
                }

            # 1. Identify contact info and name
            contact_info = []
            name = None
            name_line_idx = None
            max_contact_scan = min(8, len(lines))  # Only scan first 8 lines for contact info and name

            for idx in range(max_contact_scan):
                line = lines[idx]
                # Check for contact info
                is_email = EMAIL_PATTERN.search(line)
                is_phone = PHONE_PATTERN.search(line)
                is_location = any(loc in line.lower() for loc in LOCATION_KEYWORDS)
                if is_email or is_phone or is_location:
                    contact_info.append(line)
                elif not name and len(line.split()) <= 6 and not any(x in line for x in ['@', '|', 'www.', '.com']):
                    # Heuristic: first non-contact, non-weird line with <= 6 words is likely the name
                    name = line
                    name_line_idx = idx

            # If name not found, fallback to first line
            if not name:
                name = lines[0]
                name_line_idx = 0

            # Remove name and contact lines from the main content
            skip_idxs = set()
            if name_line_idx is not None:
                skip_idxs.add(name_line_idx)
            for idx, line in enumerate(lines[:max_contact_scan]):
                if line in contact_info:
                    skip_idxs.add(idx)
            content_lines = [line for idx, line in enumerate(lines) if idx not in skip_idxs]

            # 2. Parse sections
            sections = {
                'header': name,
                'contact': "\n".join(contact_info),
                'summary': '',
                'experience': '',
                'education': '',
                'skills': '',
                'other': ''
            }
            current_section = None
            section_buffers = {
                'summary': [],
                'experience': [],
                'education': [],
                'skills': [],
                'other': []
            }

            def match_section(line):
                for sec, pat in SECTION_HEADERS.items():
                    if pat.match(line):
                        return sec
                return None

            # Track which sections have been seen
            seen_sections = set()
            for line in content_lines:
                sec = match_section(line)
                if sec:
                    current_section = sec
                    seen_sections.add(sec)
                    continue
                # If line is a section header not recognized, treat as "other"
                if re.match(r"^[A-Z][A-Za-z\s]+$", line) and len(line.split()) <= 4 and not sec:
                    current_section = "other"
                    continue
                # Assign line to current section, or to "other" if none
                if current_section:
                    section_buffers[current_section].append(line)
                else:
                    section_buffers["other"].append(line)

            # 3. Populate sections with joined text
            for sec in ['summary', 'experience', 'education', 'skills', 'other']:
                sections[sec] = "\n".join(section_buffers[sec]).strip()

            # If all main sections are empty, put everything in experience as fallback
            if not any(sections[sec] for sec in ['experience', 'education', 'skills', 'summary']):
                sections['experience'] = "\n".join(content_lines).strip()
                sections['other'] = ''

            print(f"📝 Parsed resume: name='{sections['header']}', contact_lines={len(contact_info)}")
            for sec in ['summary', 'experience', 'education', 'skills', 'other']:
                print(f"📝 Section '{sec}': {len(sections[sec])} chars")

            return sections

        except Exception as e:
            print(f"❌ Resume parsing error: {e}")
            import traceback
            print(f"🔍 Full parsing error: {traceback.format_exc()}")
            # Safe fallback - return expected structure with all content in experience section
            return {
                'header': 'Professional Resume',
                'contact': '',
                'summary': '',
                'experience': resume_text,
                'education': '',
                'skills': '',
                'other': ''
            }
    def _format_modern_template(self, sections: Dict[str, Any]) -> str:
        """Format resume with Modern template styling and proper section formatting"""
        formatted = []

        # Header - Modern style with visual emphasis
        header_text = sections.get('header', sections.get('name', 'Professional Resume'))
        if header_text and header_text.strip():
            formatted.append("═══════════════════════════════════════════════════════")
            # Extract just the name if header contains full contact line
            if '•' in header_text:
                name_part = header_text.split('•')[0].strip()
            else:
                name_part = header_text[:50] + ('...' if len(header_text) > 50 else '')
            formatted.append(f"🔹 {name_part.upper()} 🔹")
            formatted.append("═══════════════════════════════════════════════════════")
            formatted.append("")

        # Summary section (as a paragraph)
        summary = sections.get('summary', '')
        if summary and summary.strip():
            formatted.append("📝 SUMMARY")
            formatted.append("─" * 20)
            summary_lines = [line.strip() for line in summary.split('\n') if line.strip()]
            formatted.append(" ".join(summary_lines))
            formatted.append("")

        # Experience section (bullet points, proper spacing)
        experience = sections.get('experience', '')
        if experience and experience.strip():
            formatted.append("💼 PROFESSIONAL EXPERIENCE")
            formatted.append("─" * 35)
            exp_lines = [line.strip() for line in experience.split('\n') if line.strip()]
            # Try to detect if already bulleted, otherwise add bullets
            for line in exp_lines:
                if line.startswith(("-", "•", "*")):
                    formatted.append(f"{line}")
                else:
                    formatted.append(f"• {line}")
            formatted.append("")

        # Skills section (comma-separated or categorized)
        skills = sections.get('skills', '')
        if skills and skills.strip():
            formatted.append("⚙️ TECHNICAL SKILLS")
            formatted.append("─" * 25)
            # Try to split by comma or newlines, then join as comma-separated
            skill_items = []
            for line in skills.split('\n'):
                skill_items.extend([s.strip() for s in line.split(',') if s.strip()])
            if skill_items:
                formatted.append(", ".join(skill_items))
            formatted.append("")

        # Education section (formatted entries)
        education = sections.get('education', '')
        if education and education.strip():
            formatted.append("🎓 EDUCATION")
            formatted.append("─" * 15)
            edu_lines = [line.strip() for line in education.split('\n') if line.strip()]
            for line in edu_lines:
                formatted.append(f"• {line}")
            formatted.append("")

        # Additional Information/Other section
        other = sections.get('other', '')
        if other and other.strip():
            formatted.append("📄 ADDITIONAL INFORMATION")
            formatted.append("─" * 30)
            other_lines = [line.strip() for line in other.split('\n') if line.strip()]
            for line in other_lines:
                formatted.append(f"• {line}")

        # If no content at all, fallback to a minimal message
        if not any([
            summary and summary.strip(),
            experience and experience.strip(),
            skills and skills.strip(),
            education and education.strip(),
            other and other.strip()
        ]):
            formatted.append("No resume content available.")

        final_result = '\n'.join(formatted)
        return final_result
    def _format_classic_template(self, sections: Dict[str, Any]) -> str:
        """Format resume with Classic template styling"""
        formatted = []
        
        # Header - Classic style
        header_text = sections.get('header', sections.get('name', 'Professional Resume'))
        if header_text:
            # Extract name part for classic header
            if '•' in header_text:
                name_part = header_text.split('•')[0].strip()
            else:
                name_part = header_text[:50] + ('...' if len(header_text) > 50 else '')
            formatted.append(name_part.upper())
            formatted.append("=" * len(name_part))
            formatted.append("")
        
        # Experience section
        experience = sections.get('experience', '')
        if experience:
            formatted.append("PROFESSIONAL EXPERIENCE")
            formatted.append("-" * 23)
            formatted.append(experience)
            formatted.append("")
        
        # Skills section
        skills = sections.get('skills', '')
        if skills:
            formatted.append("SKILLS & COMPETENCIES")
            formatted.append("-" * 21)
            formatted.append(skills)
            formatted.append("")
        
        # Education section
        education = sections.get('education', '')
        if education:
            formatted.append("EDUCATION")
            formatted.append("-" * 9)
            formatted.append(education)
            formatted.append("")
        
        # Other content
        other = sections.get('other', '')
        if other:
            formatted.append("ADDITIONAL INFORMATION")
            formatted.append("-" * 22)
            formatted.append(other)
        
        return '\n'.join(formatted)
    
    def _format_technical_template(self, sections: Dict[str, Any]) -> str:
        """Format resume with Technical template styling"""
        formatted = []
        
        # Header - Technical style with code-like formatting
        header_text = sections.get('header', sections.get('name', 'Professional Resume'))
        if header_text:
            # Extract name part for technical header
            if '•' in header_text:
                name_part = header_text.split('•')[0].strip()
            else:
                name_part = header_text[:50] + ('...' if len(header_text) > 50 else '')
            formatted.append(f"/* ========================= */")
            formatted.append(f"/*   {name_part.upper()}   */")
            formatted.append(f"/* ========================= */")
            formatted.append("")
        
        # Experience section - Technical style
        experience = sections.get('experience', '')
        if experience:
            formatted.append("// ====================")
            formatted.append("// PROFESSIONAL EXPERIENCE")
            formatted.append("// ====================")
            formatted.append(experience)
            formatted.append("")
        
        # Skills section - Technical style
        skills = sections.get('skills', '')
        if skills:
            formatted.append("// ===============")
            formatted.append("// TECHNICAL STACK")
            formatted.append("// ===============")
            formatted.append(skills)
            formatted.append("")
        
        # Education section
        education = sections.get('education', '')
        if education:
            formatted.append("// ===========")
            formatted.append("// EDUCATION")
            formatted.append("// ===========")
            formatted.append(education)
            formatted.append("")
        
        # Other content
        other = sections.get('other', '')
        if other:
            formatted.append("// ====================")
            formatted.append("// ADDITIONAL INFO")
            formatted.append("// ====================")
            formatted.append(other)
        
        return '\n'.join(formatted)
    
    def generate_professional_docx(self, resume_text: str, template: str = "modern") -> Dict[str, Any]:
        """Generate professional DOCX with template styling"""
        try:
            if not DEPENDENCIES_AVAILABLE:
                return {'success': False, 'error': 'DOCX dependencies not available'}
            
            # Create document
            doc = Document()
            
            # Apply template styling
            self._apply_docx_template(doc, resume_text, template)
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return {
                'success': True,
                'docx_content': docx_buffer.getvalue(),
                'template_used': template
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _apply_ats_optimizations(self, resume_text: str, ats_results: Dict[str, Any]) -> str:
        """Apply ATS optimizations based on scoring results"""
        optimized_text = resume_text
        
        # Fix bullet text issues
        optimized_text = re.sub(r'\\bullet', '•', optimized_text)
        
        # Improve section headers
        optimized_text = re.sub(r'^([A-Z\s]+):?\s*$', r'\1', optimized_text, flags=re.MULTILINE)
        
        # Add spacing for better structure
        optimized_text = re.sub(r'\n([A-Z][A-Z\s]+)\n', r'\n\n\1\n', optimized_text)
        
        return optimized_text
    
    def _apply_docx_template(self, doc: Document, resume_text: str, template: str):
        """Apply template styling to DOCX document"""
        # Parse resume sections
        # Sanitize first to remove junk panels
        try:
            sanitize_input_text = _prefer_backend('backend.services.content_filters', 'services.content_filters', attr='sanitize_input_text')
        except Exception:
            sanitize_input_text = lambda x: x
        sections = self._parse_resume_sections(sanitize_input_text(resume_text))
        
        # Apply template-specific formatting
        for section_name, content in sections.items():
            if section_name == 'header':
                self._add_docx_header(doc, content, template)
            else:
                self._add_docx_section(doc, section_name, content, template)
    
    def _parse_resume_sections(self, text: str) -> Dict[str, List[str]]:
        """Parse resume into sections"""
        sections = {'header': [], 'experience': [], 'education': [], 'skills': [], 'other': []}
        current_section = 'header'
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Detect section headers
            if any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'summary']):
                if 'experience' in line.lower():
                    current_section = 'experience'
                elif 'education' in line.lower():
                    current_section = 'education'
                elif 'skills' in line.lower():
                    current_section = 'skills'
                else:
                    current_section = 'other'
                continue
            
            sections[current_section].append(line)
        
        return sections
    
    def _add_docx_header(self, doc: Document, content: List[str], template: str):
        """Add header section to DOCX"""
        if content:
            # Name (first line, larger font)
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(content[0])
            name_run.font.size = Pt(18)
            name_run.font.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            for line in content[1:]:
                contact_para = doc.add_paragraph()
                contact_run = contact_para.add_run(line)
                contact_run.font.size = Pt(10)
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_docx_section(self, doc: Document, section_name: str, content: List[str], template: str):
        """Add content section to DOCX"""
        if not content:
            return
        
        # Section header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(section_name.upper())
        header_run.font.size = Pt(12)
        header_run.font.bold = True
        
        # Section content
        for line in content:
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(line)
            content_run.font.size = Pt(10)
    
    def get_available_templates(self) -> List[Dict[str, str]]:
        """Get list of available templates (restricted to executive_compact)."""
        return [
            {'name': 'executive_compact', 'description': 'Executive compact one-page (default)'}
        ]

    def export_scorecard(self, ats_score: Dict[str, Any], fmt: str = "json") -> Dict[str, Any]:
        """Emit ATS scorecard as a separate artifact (JSON for now)."""
        try:
            if fmt.lower() == "json":
                import json
                return {
                    'success': True,
                    'mime': 'application/json',
                    'content': json.dumps(ats_score).encode('utf-8'),
                    'filename': 'scorecard.json',
                }
            # Future: support PDF rendering for scorecard if needed
            return {'success': False, 'error': 'Unsupported format', 'mime': 'text/plain'}
        except Exception as e:
            return {'success': False, 'error': str(e), 'mime': 'text/plain'}


def make_short_preview_string(resume_json: Dict[str, Any]) -> str:
    """Create a concise preview string for UI cards."""
    try:
        name = (resume_json.get('name') or '').strip()
        headline = (resume_json.get('headline') or '').strip()
        first_skill = ''
        skills = resume_json.get('skills') or []
        if isinstance(skills, list) and skills:
            first = skills[0]
            if isinstance(first, dict):
                first_skill = first.get('name') or ''
            elif isinstance(first, str):
                first_skill = first
        first_exp = ''
        experience = resume_json.get('experience') or []
        if isinstance(experience, list) and experience:
            item = experience[0]
            if isinstance(item, dict):
                company = (item.get('company') or '').strip()
                title = (item.get('title') or '').strip()
                first_exp = f"{company} ({title})" if company or title else ''
        head_or_skill = headline or first_skill
        pieces = [p for p in [name, head_or_skill, first_exp] if p]
        return ' — '.join(pieces)
    except Exception:
        return (resume_json.get('name') or '').strip()
